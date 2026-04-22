#!/usr/bin/env python3
"""
Document Converter Backend Server
파일 형식별 HTML 변환 백엔드 서비스

Usage:
  pip3 install flask python-docx openpyxl pdfplumber
  python3 convert-server.py [port]
  Default port: 8082

API:
  POST /api/convert
    Content-Type: multipart/form-data
    Body: file=<binary>
    Response: { html, text, metadata: { format, pages, sheets, warnings } }

  GET /health
    Response: { status: 'ok', formats: [...] }
"""

import sys
import os
import json
import tempfile
import subprocess
import zipfile
import re
import struct
import zlib
from io import BytesIO
from html import escape as html_escape

try:
    from flask import Flask, request, jsonify

    FLASK_AVAILABLE = True
except ImportError:
    FLASK_AVAILABLE = False

# ============================================
# Flask App
# ============================================

if FLASK_AVAILABLE:
    app = Flask(__name__)

    # CORS headers
    @app.after_request
    def add_cors(response):
        response.headers["Access-Control-Allow-Origin"] = "*"
        response.headers["Access-Control-Allow-Methods"] = "GET, POST, OPTIONS"
        response.headers["Access-Control-Allow-Headers"] = "*"
        return response

    @app.route("/health", methods=["GET"])
    def health():
        formats = get_supported_formats()
        return jsonify(
            {
                "status": "ok",
                "version": "build75",
                "formats": formats,
                "libreoffice": SOFFICE_PATH or False,
                "hwp_table_support": bool(SOFFICE_PATH),
                "playwright": check_playwright(),
            }
        )

    @app.route("/api/render-full", methods=["POST", "OPTIONS"])
    def render_full():
        """Preview 전체 HTML → PNG (Playwright full_page 스크린샷)

        Request JSON:
          { html: string, width?: int, scale?: float, format?: 'png'|'jpeg' }
        Response:
          image/png 또는 image/jpeg 바이너리
        """
        if request.method == "OPTIONS":
            return "", 204

        payload = request.get_json(silent=True) or {}
        img_bytes, mime, err = _render_full_playwright(payload)
        if err:
            status = 503 if "설치" in err else 500
            return jsonify({"error": err}), status

        from flask import Response
        return Response(img_bytes, mimetype=mime)

    @app.route("/api/render-kv", methods=["POST", "OPTIONS"])
    def render_kv():
        """KV HTML → PNG 이미지 변환 (Playwright 서버 사이드 렌더링)

        Request JSON:
          { html: string, width: int, height: int, scale: float, format: 'png'|'jpeg' }
        Response:
          image/png 또는 image/jpeg 바이너리
        """
        if request.method == "OPTIONS":
            return "", 204

        payload = request.get_json(silent=True) or {}
        img_bytes, mime, err = _render_kv_playwright(payload)
        if err:
            status = 503 if "설치" in err else 500
            return jsonify({"error": err}), status

        from flask import Response
        return Response(img_bytes, mimetype=mime)

    @app.route("/api/convert", methods=["POST", "OPTIONS"])
    def convert():
        if request.method == "OPTIONS":
            return "", 204

        if "file" not in request.files:
            return jsonify({"error": "No file provided"}), 400

        file = request.files["file"]
        if not file.filename:
            return jsonify({"error": "Empty filename"}), 400

        ext = file.filename.rsplit(".", 1)[-1].lower()
        data = file.read()

        try:
            parser_hint = request.args.get("parser", "")
            api_key = request.form.get('anthropic_api_key', '') or os.environ.get('ANTHROPIC_API_KEY', '')
            result = convert_file(data, ext, file.filename, parser_hint=parser_hint, api_key=api_key)
            return jsonify(result)
        except Exception as e:
            return (
                jsonify(
                    {
                        "error": str(e),
                        "html": "",
                        "text": "",
                        "metadata": {"format": ext, "warnings": [str(e)]},
                    }
                ),
                500,
            )

    @app.route("/api/hwp-to-pdf", methods=["POST", "OPTIONS"])
    def hwp_to_pdf():
        """HWP/HWPX → PDF 변환 (Vision 변환용). base64 PDF 반환."""
        if request.method == "OPTIONS":
            return "", 204

        if "file" not in request.files:
            return jsonify({"error": "No file provided"}), 400

        file = request.files["file"]
        if not file.filename:
            return jsonify({"error": "Empty filename"}), 400

        ext = file.filename.rsplit(".", 1)[-1].lower()
        if ext not in ("hwp", "hwpx", "docx", "doc", "pdf"):
            return (
                jsonify({"error": f".{ext} 파일은 Vision 변환을 지원하지 않습니다."}),
                400,
            )

        data = file.read()

        try:
            import base64

            if ext == "pdf":
                # 이미 PDF면 그대로 반환
                pdf_b64 = base64.b64encode(data).decode()
                return jsonify({"pdf": pdf_b64, "pages": None, "warnings": []})

            pdf_data, error = convert_any_to_pdf(data, ext, file.filename)
            if error:
                return jsonify({"error": error}), 500

            pdf_b64 = base64.b64encode(pdf_data).decode()
            return jsonify({"pdf": pdf_b64, "warnings": []})
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    @app.route("/api/hwp-to-rawhtml", methods=["POST", "OPTIONS"])
    def hwp_to_rawhtml():
        """HWP/HWPX → 원본 HTML 변환 (AI 처리 전 단계용). LibreOffice HTML 직접 출력."""
        if request.method == "OPTIONS":
            return "", 204

        if "file" not in request.files:
            return jsonify({"error": "No file provided"}), 400

        file = request.files["file"]
        if not file.filename:
            return jsonify({"error": "Empty filename"}), 400

        ext = file.filename.rsplit(".", 1)[-1].lower()
        if ext not in ("hwp", "hwpx", "docx", "doc"):
            return (
                jsonify(
                    {"error": f".{ext} 파일은 원문 HTML 변환을 지원하지 않습니다."}
                ),
                400,
            )

        data = file.read()
        try:
            html, warnings = convert_to_raw_html(data, ext)
            if not html:
                return (
                    jsonify(
                        {
                            "error": "변환 실패: HTML을 생성하지 못했습니다.",
                            "warnings": warnings,
                        }
                    ),
                    500,
                )
            return jsonify({"html": html, "warnings": warnings})
        except Exception as e:
            return jsonify({"error": str(e)}), 500


# ============================================
# Format Detection & Routing
# ============================================


_playwright_available = None  # None = 미확인, True/False = 확인됨

def check_playwright():
    """Playwright + Chromium 설치 여부 확인 (결과 캐싱, 최초 1회만 실행)."""
    global _playwright_available
    if _playwright_available is not None:
        return _playwright_available
    try:
        from playwright.sync_api import sync_playwright
        with sync_playwright() as p:
            browser = p.chromium.launch(args=["--no-sandbox"])
            browser.close()
        _playwright_available = True
    except Exception:
        _playwright_available = False
    return _playwright_available


def _render_full_playwright(payload):
    """Preview 전체 HTML → PNG 바이트 (Playwright full_page 스크린샷).

    Request payload:
      html (str, 필수): 전체 HTML (외부 <link>, <style>, 본문 포함)
      width (int, 기본 900): viewport 폭
      scale (float, 기본 2.0): device pixel ratio
      format ('png'|'jpeg', 기본 'png'): 출력 포맷
    Returns: (bytes, mime_type, error_str)
    """
    try:
        from playwright.sync_api import sync_playwright
    except ImportError:
        return None, None, "Playwright가 설치되지 않았습니다. pip install playwright && playwright install chromium"

    html_content = payload.get("html", "")
    width = int(payload.get("width", 900))
    scale = float(payload.get("scale", 2.0))
    fmt = payload.get("format", "png").lower()
    if fmt not in ("png", "jpeg", "pdf"):
        fmt = "png"
    if not html_content:
        return None, None, "html 파라미터가 필요합니다"

    # body/html 기본 여백 제거 + 콘텐츠를 900px 폭으로 가운데 정렬 없이 꽉차게
    reset_css = (
        "<style>"
        "html,body{margin:0!important;padding:0!important;}"
        "body{background:#ffffff!important;}"
        "#templwrap_v3,#templwrap{margin:0!important;}"
        "</style>"
    )
    # <head> 내부에 reset_css 주입 (없으면 맨 앞에 prepend)
    if "<head" in html_content.lower():
        # <head ...>\n뒤에 삽입 — 대소문자 고려
        import re as _re
        html_content = _re.sub(r"(<head[^>]*>)", r"\1" + reset_css, html_content, count=1, flags=_re.IGNORECASE)
    else:
        html_content = reset_css + html_content

    try:
        with sync_playwright() as p:
            browser = p.chromium.launch(args=["--no-sandbox", "--disable-setuid-sandbox"])
            page = browser.new_page(
                viewport={"width": width, "height": 800},
                device_scale_factor=scale,
            )
            page.set_content(html_content, wait_until="networkidle")
            try:
                page.wait_for_load_state("networkidle", timeout=5000)
            except Exception:
                pass
            if fmt == "pdf":
                # PDF — 콘텐츠 폭과 동일하게 맞추고 높이는 자연 흐름대로
                # print_background=True로 배경 유지. 여백 0으로 원본 그대로.
                pdf_bytes = page.pdf(
                    width=f"{width}px",
                    print_background=True,
                    prefer_css_page_size=False,
                    margin={"top": "0", "right": "0", "bottom": "0", "left": "0"},
                )
                browser.close()
                return pdf_bytes, "application/pdf", None
            # PNG / JPEG
            screenshot_bytes = page.screenshot(type=fmt, full_page=True)
            browser.close()

        mime = "image/jpeg" if fmt == "jpeg" else "image/png"
        return screenshot_bytes, mime, None
    except Exception as e:
        return None, None, f"렌더링 오류: {str(e)}"


def _render_kv_playwright(payload):
    """KV HTML → PNG/JPEG 바이트 변환 (Flask/stdlib 공유 로직).

    Returns: (bytes, mime_type, error_str)
    """
    try:
        from playwright.sync_api import sync_playwright
    except ImportError:
        return None, None, "Playwright가 설치되지 않았습니다. pip install playwright && playwright install chromium"

    html_content = payload.get("html", "")
    width = int(payload.get("width", 1200))
    height = int(payload.get("height", 675))
    scale = float(payload.get("scale", 2.0))
    fmt = payload.get("format", "png").lower()
    if fmt not in ("png", "jpeg"):
        fmt = "png"
    if not html_content:
        return None, None, "html 파라미터가 필요합니다"

    try:
        with sync_playwright() as p:
            browser = p.chromium.launch(args=["--no-sandbox", "--disable-setuid-sandbox"])
            page = browser.new_page(
                viewport={"width": width, "height": height},
                device_scale_factor=scale,
            )
            page.set_content(html_content, wait_until="networkidle")
            try:
                page.wait_for_load_state("networkidle", timeout=3000)
            except Exception:
                pass
            screenshot_bytes = page.screenshot(
                type=fmt,
                clip={"x": 0, "y": 0, "width": width, "height": height},
                full_page=False,
            )
            browser.close()

        mime = "image/jpeg" if fmt == "jpeg" else "image/png"
        return screenshot_bytes, mime, None
    except Exception as e:
        return None, None, f"렌더링 오류: {str(e)}"


def get_supported_formats():
    formats = ["txt", "html", "htm", "csv"]
    try:
        import docx

        formats.extend(["docx"])
    except ImportError:
        pass
    try:
        import openpyxl

        formats.extend(["xlsx", "xls"])
    except ImportError:
        pass
    try:
        import pdfplumber

        formats.append("pdf")
    except ImportError:
        pass
    # HWP/HWPX — always attempt
    formats.extend(["hwp", "hwpx", "doc"])
    return list(set(formats))


def convert_file(data, ext, filename="", parser_hint="", api_key=""):
    """Route to format-specific converter."""
    converters = {
        "docx": convert_docx,
        "doc": convert_doc,
        "xlsx": convert_xlsx,
        "xls": convert_xlsx,
        "pdf": convert_pdf,
        "hwp": convert_hwp,
        "hwpx": convert_hwpx,
        "txt": convert_text,
        "csv": convert_csv,
        "html": convert_html_file,
        "htm": convert_html_file,
    }

    converter = converters.get(ext)
    if not converter:
        raise ValueError(f"지원하지 않는 파일 형식: .{ext}")

    if ext == "hwp":
        result = converter(data, filename, parser_hint=parser_hint, api_key=api_key)
    else:
        result = converter(data, filename)
    result.setdefault("metadata", {})
    result["metadata"]["format"] = ext
    result["metadata"]["filename"] = filename
    return result


# ============================================
# DOCX Converter (python-docx)
# ============================================


# ============================================
# 한컴 PUA(Private Use Area) 문자 정리
# HWP → DOCX 변환 시 한컴 폰트의 특수문자가 PUA(U+E000~U+F8FF)로 깨지는 문제 처리
# ============================================
HWP_PUA_MAP = {
    # 가장 흔한 깨진 중점/구분자 (한컴 폰트 PUA)
    '\uF06C': '·', '\uF0B7': '·', '\uF09F': '·', '\uF0FB': '·',
    '\uE0FC': '·', '\uE0FB': '·', '\uF076': '·',
    # 화살표
    '\uF0E0': '→', '\uF0E1': '↑', '\uF0E2': '↓', '\uF0DF': '←',
    # 체크/별
    '\uF0FC': '✓', '\uF0FE': '☑', '\uF0AB': '★', '\uF0AC': '☆',
    # 기호
    '\uF0A7': '※', '\uF0B0': '°', '\uF0B1': '±', '\uF0B2': '²', '\uF0B3': '³',
    '\uF0BC': '¼', '\uF0BD': '½', '\uF0BE': '¾',
    # 따옴표
    '\uF022': '"', '\uF027': "'",
}

def _clean_hwp_pua(text):
    """HWP 변환 결과에서 PUA 깨짐 문자를 일반 유니코드로 복원.
    매핑이 없는 PUA 문자는 한글 사이일 때 중점(·)으로 추정 대체."""
    if not text or not isinstance(text, str):
        return text
    # 1) 알려진 매핑 적용
    for pua, replacement in HWP_PUA_MAP.items():
        if pua in text:
            text = text.replace(pua, replacement)
    # 2) 남은 PUA 문자 처리: 한글 사이면 중점, 아니면 제거
    import re
    def _replace_pua(m):
        ch = m.group(0)
        idx = m.start()
        prev_ch = text[idx-1] if idx > 0 else ''
        next_ch = text[idx+1] if idx+1 < len(text) else ''
        # 한글 사이면 중점으로 추정
        if re.match(r'[가-힣]', prev_ch) and re.match(r'[가-힣]', next_ch):
            return '·'
        return ''  # 그 외는 제거
    text = re.sub(r'[\uE000-\uF8FF]', _replace_pua, text)
    return text


def convert_docx(data, filename=""):
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.text.paragraph import Paragraph as DocxParagraph
        from docx.table import Table as DocxTable
    except ImportError:
        raise ImportError("python-docx가 설치되지 않았습니다. pip3 install python-docx")

    doc = Document(BytesIO(data))
    html_parts = []
    warnings = []
    bold_texts = []    # run.bold == True 인 텍스트 목록
    underline_texts = []  # run.underline == True 인 텍스트 목록
    bullet_items = []  # {"bullet": "·", "text": "..."} 목록 (단락 전체 텍스트)

    # WML 네임스페이스 (OOXML)
    WML_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    def _effective_bold(run):
        """run.bold가 None(스타일 상속)인 경우 단락/문서 기본 스타일까지 탐색"""
        if run.bold is True:
            return True
        if run.bold is False:
            return False
        # None → 단락 스타일 체크
        try:
            para = run._element.getparent()
            # pPr → rPr → b 체크 (단락 기본 런 서식)
            ppr = para.find(f"{WML_NS}pPr")
            if ppr is not None:
                rpr = ppr.find(f"{WML_NS}rPr")
                if rpr is not None:
                    b_elem = rpr.find(f"{WML_NS}b")
                    if b_elem is not None:
                        val = b_elem.get(f"{WML_NS}val", "true")
                        return val.lower() not in ("false", "0")
        except Exception:
            pass
        return False

    def _effective_underline(run):
        """run.underline이 None인 경우 단락 스타일까지 탐색"""
        if run.underline is True or (run.underline and run.underline is not False):
            return True
        if run.underline is False:
            return False
        try:
            para = run._element.getparent()
            ppr = para.find(f"{WML_NS}pPr")
            if ppr is not None:
                rpr = ppr.find(f"{WML_NS}rPr")
                if rpr is not None:
                    u_elem = rpr.find(f"{WML_NS}u")
                    if u_elem is not None:
                        val = u_elem.get(f"{WML_NS}val", "single")
                        return val.lower() not in ("none", "false", "0")
        except Exception:
            pass
        return False

    def _run_to_html(run):
        """run → bold/italic/underline + 폰트 크기/색상 inline style HTML"""
        t = html_escape(run.text)
        if not t:
            return ""
        run_styles = []
        rpr = run._element.find(f"{WML_NS}rPr")
        if rpr is not None:
            # 폰트 크기 (sz = half-points)
            for sz_tag in [f"{WML_NS}sz", f"{WML_NS}szCs"]:
                sz_elem = rpr.find(sz_tag)
                if sz_elem is not None:
                    try:
                        sz_val = int(
                            sz_elem.get(f"{WML_NS}val", sz_elem.get("val", "0"))
                        )
                        if sz_val > 0:
                            run_styles.append(f"font-size:{sz_val // 2}pt")
                    except Exception:
                        pass
                    break
            # 글자색
            color_elem = rpr.find(f"{WML_NS}color")
            if color_elem is not None:
                color_val = color_elem.get(f"{WML_NS}val", color_elem.get("val", ""))
                if color_val and color_val.upper() not in ("AUTO", "FFFFFF", ""):
                    run_styles.append(f"color:#{color_val}")
        if _effective_bold(run):
            raw = run.text.strip()
            if raw and len(raw) <= 80:
                bold_texts.append(raw)
            t = f"<strong>{t}</strong>"
        if run.italic:
            t = f"<em>{t}</em>"
        if _effective_underline(run):
            raw = run.text.strip()
            if raw and len(raw) <= 80:
                underline_texts.append(raw)
            t = f"<u>{t}</u>"
        if run_styles:
            t = f'<span style="{";".join(run_styles)}">{t}</span>'
        return t

    def _para_style(para):
        """단락 정렬·들여쓰기·줄높이 → CSS style 문자열"""
        styles = []
        ppr = para._element.pPr
        if ppr is None:
            return ""
        # 수평 정렬
        jc_elem = ppr.find(f"{WML_NS}jc")
        if jc_elem is not None:
            jc_val = jc_elem.get(f"{WML_NS}val", jc_elem.get("val", ""))
            align_map = {
                "center": "center",
                "right": "right",
                "both": "justify",
                "distribute": "justify",
            }
            if jc_val in align_map:
                styles.append(f"text-align:{align_map[jc_val]}")
        # 들여쓰기 제거 (불필요한 padding-left, text-indent 스타일 생성 방지)
        # 줄 높이 (240 = 1.0배, 360 = 1.5배)
        spacing_elem = ppr.find(f"{WML_NS}spacing")
        if spacing_elem is not None:
            line_val = spacing_elem.get(f"{WML_NS}line", spacing_elem.get("line", ""))
            line_rule = spacing_elem.get(
                f"{WML_NS}lineRule", spacing_elem.get("lineRule", "")
            )
            if line_val and line_rule == "auto":
                try:
                    lh = int(line_val) / 240
                    if abs(lh - 1.0) > 0.1:
                        styles.append(f"line-height:{lh:.2f}")
                except Exception:
                    pass
        return ";".join(styles)

    def _tc_style(tc_pr):
        """tcPr → 셀 vertical-align·padding·background-color·width CSS 문자열"""
        if tc_pr is None:
            return ""
        styles = []
        # 수직 정렬
        va_elem = tc_pr.find(f"{WML_NS}vAlign")
        if va_elem is not None:
            va_val = va_elem.get(f"{WML_NS}val", va_elem.get("val", ""))
            va_map = {"top": "top", "center": "middle", "bottom": "bottom"}
            if va_val in va_map:
                styles.append(f"vertical-align:{va_map[va_val]}")
        # 셀 여백 (tcMar, twips → px)
        tc_mar = tc_pr.find(f"{WML_NS}tcMar")
        if tc_mar is not None:
            pads = []
            for side in ["top", "left", "bottom", "right"]:
                side_elem = tc_mar.find(f"{WML_NS}{side}")
                if side_elem is not None:
                    try:
                        w_val = int(side_elem.get(f"{WML_NS}w", "0"))
                        pads.append(f"{w_val * 96 / 1440:.0f}px")
                    except Exception:
                        pads.append("0")
                else:
                    pads.append("0")
            if any(p != "0" for p in pads):
                styles.append(f'padding:{" ".join(pads)}')
        # 배경색 (shd)
        shd_elem = tc_pr.find(f"{WML_NS}shd")
        if shd_elem is not None:
            fill = shd_elem.get(f"{WML_NS}fill", shd_elem.get("fill", ""))
            if fill and fill.upper() not in ("AUTO", "FFFFFF", ""):
                styles.append(f"background-color:#{fill}")
        # 셀 너비 (tcW, twips → px 또는 %)
        tcw_elem = tc_pr.find(f"{WML_NS}tcW")
        if tcw_elem is not None:
            w_type = tcw_elem.get(f"{WML_NS}type", tcw_elem.get("type", ""))
            w_val_str = tcw_elem.get(f"{WML_NS}w", tcw_elem.get("w", ""))
            if w_val_str:
                try:
                    if w_type == "dxa":
                        px = int(w_val_str) * 96 / 1440
                        if px > 0:
                            styles.append(f"width:{px:.0f}px")
                    elif w_type == "pct":
                        pct = int(w_val_str) / 50
                        if pct > 0:
                            styles.append(f"width:{pct:.1f}%")
                except Exception:
                    pass
        return ";".join(styles)

    def _cell_to_html(cell):
        """셀 내부 단락 → 인라인 서식 + 단락 정렬 포함 HTML (<br> 연결)"""
        parts = []
        for para in cell.paragraphs:
            runs_html = "".join(_run_to_html(run) for run in para.runs)
            p_style = _para_style(para)
            if p_style and runs_html:
                parts.append(f'<span style="{p_style}">{runs_html}</span>')
            else:
                # 빈 단락도 유지해야 줄바꿈이 살아남
                parts.append(runs_html if runs_html else '')
        # 후행 빈 단락만 제거 (단, 연속된 줄바꿈은 유지)
        while parts and not parts[-1].strip():
            parts.pop()
        # 빈 문자열도 포함해서 <br>로 연결 (줄바꿈 보존)
        return "<br>".join(parts) if parts else ''

    def _get_list_bullet(para):
        """DOCX numPr → 실제 불릿 문자 또는 번호 형식 힌트 반환.
        HWP→DOCX 변환 시 불릿 기호가 lvlText에 보존됨."""
        try:
            ppr = para._element.pPr
            if ppr is None:
                return None
            numpr_elem = ppr.find(f"{WML_NS}numPr")
            if numpr_elem is None:
                return None
            num_id_elem = numpr_elem.find(f"{WML_NS}numId")
            ilvl_elem = numpr_elem.find(f"{WML_NS}ilvl")
            if num_id_elem is None or ilvl_elem is None:
                return None
            num_id = int(num_id_elem.get(f"{WML_NS}val", num_id_elem.get("val", "0")))
            ilvl = int(ilvl_elem.get(f"{WML_NS}val", ilvl_elem.get("val", "0")))
            if num_id == 0:
                return None

            # numbering.xml에서 추상 번호 정의 탐색
            num_part = doc.part.numbering_part
            num_root = num_part._element

            # numId → abstractNumId 매핑
            abstract_num_id = None
            for num_el in num_root.findall(f"{WML_NS}num"):
                if num_el.get(f"{WML_NS}numId") == str(num_id):
                    abs_ref = num_el.find(f"{WML_NS}abstractNumId")
                    if abs_ref is not None:
                        abstract_num_id = abs_ref.get(f"{WML_NS}val")
                    break

            if abstract_num_id is None:
                return None

            # abstractNumId → lvl → numFmt + lvlText
            for abs_el in num_root.findall(f"{WML_NS}abstractNum"):
                if abs_el.get(f"{WML_NS}abstractNumId") == abstract_num_id:
                    for lvl_el in abs_el.findall(f"{WML_NS}lvl"):
                        if lvl_el.get(f"{WML_NS}ilvl") == str(ilvl):
                            num_fmt_el = lvl_el.find(f"{WML_NS}numFmt")
                            lvl_text_el = lvl_el.find(f"{WML_NS}lvlText")
                            num_fmt = (
                                num_fmt_el.get(f"{WML_NS}val", "")
                                if num_fmt_el is not None
                                else ""
                            )
                            lvl_text = (
                                lvl_text_el.get(f"{WML_NS}val", "")
                                if lvl_text_el is not None
                                else ""
                            )
                            # 불릿 문자 추출: num_fmt="bullet"이거나 lvl_text에 실제 기호가 있으면 반환
                            if lvl_text:
                                # lvl_text의 첫 문자가 실제 불릿 기호(○, ●, -, *, 등)인지 확인
                                first_char = lvl_text[0] if lvl_text else ""
                                # 번호/자리표시자가 아닌 실제 기호면 그대로 반환
                                if first_char not in "%0123456789" and num_fmt == "bullet":
                                    return first_char
                                # 또는 HWP→DOCX에서 lvl_text가 실제 기호인 경우 (예: "-", "○" 등)
                                if num_fmt == "bullet" and first_char in "-–—•○●■□▶★◆◇▪▣":
                                    return first_char
                            # 번호 형식 힌트 반환 (기호가 없을 때만)
                            fmt_hint = {
                                "decimal": "1.",
                                "lowerLetter": "a.",
                                "upperLetter": "A.",
                                "lowerRoman": "i.",
                                "upperRoman": "I.",
                                "koreanCounting": "가.",
                                "koreanDigital": "①",
                                "koreanLegal": "가.",
                                "decimalEnclosedCircle": "①",
                            }
                            return fmt_hint.get(num_fmt)
        except Exception:
            pass
        return None

    def _process_para(para):
        """단락 → HTML 문자열 (비어있으면 None)"""
        if not para.text.strip():
            return None
        style_name = para.style.name.lower() if para.style else ""
        if "heading 1" in style_name:
            tag = "h1"
        elif "heading 2" in style_name:
            tag = "h2"
        elif "heading 3" in style_name:
            tag = "h3"
        elif "list" in style_name:
            tag = "li"
        else:
            tag = "p"
        runs_html = "".join(_run_to_html(run) for run in para.runs)
        if not runs_html:
            return None

        # 제목(h1-h6)이면 불필요한 <strong>, <u> 제거
        if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
            runs_html = runs_html.replace('<strong>', '').replace('</strong>', '')
            runs_html = runs_html.replace('<u>', '').replace('</u>', '')

        p_style = _para_style(para)
        style_attr = f' style="{p_style}"' if p_style else ""

        # 리스트 단락이면 실제 불릿 문자를 텍스트 앞에 포함
        if tag == "li":
            bullet_char = _get_list_bullet(para)
            if bullet_char:
                para_text = para.text.strip()
                if para_text:
                    bullet_items.append({"bullet": bullet_char, "text": para_text[:80]})
                # 불릿 기호를 텍스트 앞에 추가 (HWP→DOCX 변환 시 기호가 분리됨)
                runs_html = f"{html_escape(bullet_char)} {runs_html}"
        return f"<{tag}{style_attr}>{runs_html}</{tag}>"

    def _classify_table(table):
        """ki-it.com 논문 방법론: lettered(내용 표) vs unlettered(레이아웃 표) 분류.
        - 단일 셀: unlettered
        - 1행 테이블: unlettered (LibreOffice가 HWP 텍스트박스/2단 레이아웃을 1행 다열 테이블로 변환하는 케이스)
        - 비어있는 셀 비율이 85% 초과: unlettered
        - 그 외: lettered
        """
        rows = table.rows
        num_rows = len(rows)
        if num_rows == 0:
            return "unlettered"

        total_cells = 0
        non_empty_cells = 0
        for row in rows:
            seen_tcs = set()
            for cell in row.cells:
                tc = cell._tc
                if id(tc) in seen_tcs:
                    continue
                seen_tcs.add(id(tc))
                total_cells += 1
                if cell.text.strip():
                    non_empty_cells += 1

        if total_cells <= 1:
            return "unlettered"

        # 1행 테이블: LibreOffice가 HWP 텍스트박스·2단 레이아웃을 1행 다열로 변환하는 경우
        # 헤더 행만 있는 1행 테이블도 해당됨 → 레이아웃 표로 판정
        if num_rows == 1:
            return "unlettered"

        fill_ratio = non_empty_cells / total_cells
        return "lettered" if fill_ratio >= 0.15 else "unlettered"

    def _process_table(table):
        """테이블 → HTML 줄 목록 반환 (ki-it.com: lettered/unlettered 분류 적용)"""
        if _classify_table(table) == "unlettered":
            # 레이아웃 표: 셀 내용을 <p> 단락으로 변환
            parts = []
            for row in table.rows:
                seen_tcs = set()
                for cell in row.cells:
                    tc = cell._tc
                    if id(tc) in seen_tcs:
                        continue
                    seen_tcs.add(id(tc))
                    cell_html = _cell_to_html(cell)
                    if cell_html.strip():
                        parts.append(f"<p>{cell_html}</p>")
            return parts

        # lettered: 내용 표 → <table> HTML 유지
        parts = ["<table>"]
        for i, row in enumerate(table.rows):
            parts.append("  <tr>")
            seen_tcs = set()

            for j, cell in enumerate(row.cells):
                tc = cell._tc
                if id(tc) in seen_tcs:
                    continue
                seen_tcs.add(id(tc))

                try:
                    grid_span = tc.grid_span
                    colspan = grid_span if grid_span and grid_span > 1 else 1
                except Exception:
                    colspan = 1

                tc_pr = tc.find(f"{WML_NS}tcPr")
                vmerge_elem = (
                    tc_pr.find(f"{WML_NS}vMerge") if tc_pr is not None else None
                )

                if vmerge_elem is not None:
                    vmerge_val = vmerge_elem.get(
                        f"{WML_NS}val", vmerge_elem.get("val", "continue")
                    )
                    if vmerge_val == "continue":
                        continue

                rowspan = 1
                if vmerge_elem is not None:
                    for k in range(i + 1, len(table.rows)):
                        try:
                            below_cell = table.cell(k, j)
                            below_tc_pr = below_cell._tc.find(f"{WML_NS}tcPr")
                            below_vm = (
                                below_tc_pr.find(f"{WML_NS}vMerge")
                                if below_tc_pr is not None
                                else None
                            )
                            if below_vm is not None:
                                bv = below_vm.get(
                                    f"{WML_NS}val", below_vm.get("val", "continue")
                                )
                                if bv == "continue":
                                    rowspan += 1
                                else:
                                    break
                            else:
                                break
                        except Exception:
                            break

                cell_tag = "th" if i == 0 else "td"
                cell_html = _cell_to_html(cell)
                attrs = ""
                if colspan > 1:
                    attrs += f' colspan="{colspan}"'
                if rowspan > 1:
                    attrs += f' rowspan="{rowspan}"'
                cell_style = _tc_style(tc_pr)
                if cell_style:
                    attrs += f' style="{cell_style}"'
                parts.append(f"    <{cell_tag}{attrs}>{cell_html}</{cell_tag}>")

            parts.append("  </tr>")
        parts.append("</table>")
        return parts

    # 단락·테이블 원본 순서 유지하며 순회
    P_TAG = qn("w:p")
    TBL_TAG = qn("w:tbl")
    for child in doc.element.body:
        if child.tag == P_TAG:
            para = DocxParagraph(child, doc)
            line = _process_para(para)
            if line:
                html_parts.append(line)
        elif child.tag == TBL_TAG:
            table = DocxTable(child, doc)
            html_parts.extend(_process_table(table))

    html = "\n".join(html_parts)
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    # HWP PUA 깨짐 정리 (한컴 폰트 특수문자가 PUA로 변환된 경우)
    html = _clean_hwp_pua(html)
    text = _clean_hwp_pua(text)

    return {
        "html": html,
        "text": text,
        "metadata": {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "warnings": warnings,
            "bold_texts": [_clean_hwp_pua(t) for t in dict.fromkeys(bold_texts)],
            "underline_texts": [_clean_hwp_pua(t) for t in dict.fromkeys(underline_texts)],
            "bullet_items": bullet_items,  # [{"bullet": "·", "text": "..."}, ...]
        },
    }


# ============================================
# DOC Converter (textutil on macOS)
# ============================================


def convert_doc(data, filename=""):
    warnings = []

    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
        tmp.write(data)
        tmp_path = tmp.name

    try:
        # macOS textutil
        out_path = tmp_path.replace(".doc", ".html")
        result = subprocess.run(
            ["textutil", "-convert", "html", "-output", out_path, tmp_path],
            capture_output=True,
            text=True,
            timeout=30,
        )

        if result.returncode == 0 and os.path.exists(out_path):
            with open(out_path, "r", encoding="utf-8", errors="replace") as f:
                html = f.read()
            # body 내용만 추출
            body_match = re.search(
                r"<body[^>]*>(.*?)</body>", html, re.DOTALL | re.IGNORECASE
            )
            html = body_match.group(1) if body_match else html
            text = re.sub(r"<[^>]+>", " ", html)
            text = re.sub(r"\s+", " ", text).strip()
            os.unlink(out_path)
            return {"html": html, "text": text, "metadata": {"warnings": warnings}}

        warnings.append("textutil 변환 실패. 텍스트만 추출합니다.")

    except (FileNotFoundError, subprocess.TimeoutExpired) as e:
        warnings.append(f"textutil 사용 불가: {e}")
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)

    # Fallback: 바이너리에서 텍스트 추출
    text = data.decode("utf-8", errors="replace")
    text = re.sub(r"[^\x20-\x7E\uAC00-\uD7AF\u3131-\u318E\n\r\t.,;:!?\-]", " ", text)
    text = re.sub(r"\s{3,}", "\n", text).strip()
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    html = "\n".join(f"<p>{html_escape(l)}</p>" for l in lines)

    warnings.append(
        "DOC 파일은 제한적으로 지원됩니다. DOCX로 변환 후 다시 시도해주세요."
    )
    return {"html": html, "text": "\n".join(lines), "metadata": {"warnings": warnings}}


# ============================================
# XLSX/XLS Converter (openpyxl)
# ============================================


def convert_xlsx(data, filename=""):
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise ImportError("openpyxl이 설치되지 않았습니다. pip3 install openpyxl")

    wb = load_workbook(BytesIO(data), data_only=True)
    html_parts = []
    warnings = []
    text_parts = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        if len(wb.sheetnames) > 1:
            html_parts.append(f"<h3>{html_escape(sheet_name)}</h3>")

        html_parts.append("<table>")
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            cells = [c if c is not None else "" for c in row]
            if not any(str(c).strip() for c in cells):
                continue

            html_parts.append("  <tr>")
            for cell in cells:
                tag = "th" if i == 0 else "td"
                val = html_escape(str(cell))
                html_parts.append(f"    <{tag}>{val}</{tag}>")
                text_parts.append(str(cell))
            html_parts.append("  </tr>")
        html_parts.append("</table>")

    html = "\n".join(html_parts)
    text = "\t".join(text_parts)

    return {
        "html": html,
        "text": text,
        "metadata": {"sheets": wb.sheetnames, "warnings": warnings},
    }


# ============================================
# PDF Converter (pdfplumber)
# ============================================

# 제목이 아닌 짧은 줄 패턴 (전화번호, 이메일, 날짜, URL, 숫자+단위, 목록 기호)
_HEADING_EXCLUDE_RE = re.compile(
    r"^\s*("
    r"\d{2,4}[.\-/]\d{1,2}[.\-/]\d{1,2}"  # 날짜: 2026.01.30
    r"|0\d{1,2}[.\-)\s]\d{3,4}[.\-]\d{4}"  # 전화: 02-1234-5678
    r"|\d{2,3}[.\-]\d{3,4}[.\-]\d{4}"  # 전화: 010-1234-5678
    r"|[\w.\-]+@[\w.\-]+\.\w+"  # 이메일
    r"|https?://"  # URL
    r"|\d[\d,]*\s*(명|원|만원|천원|억|개|건|%|세|년|월|일|시간|분|회|차)"  # 숫자+단위
    r"|[가-힣]{1,3}[.:)\s]\s*\d"  # 한글 라벨: 가. 1
    r"|\d+[.)]\s"  # 번호 항목: 1. 또는 1)
    r"|[○●■□▶※★◆◎▣☐☑]\s"  # 목록 기호
    r")",
    re.IGNORECASE,
)


def _is_heading_candidate(line):
    """짧은 줄이 제목인지 판별 (오탐 제외 패턴 적용)"""
    if len(line) >= 40:
        return False
    if line.endswith(".") or line.endswith(",") or line.endswith(";"):
        return False
    if _HEADING_EXCLUDE_RE.match(line):
        return False
    korean_chars = sum(1 for c in line if "\uac00" <= c <= "\ud7af")
    if korean_chars < 2 and len(line) < 10:
        return False
    return True


def _expand_multiline_rows(data):
    """셀 내 줄바꿈을 분석하여 행 분리 + rowspan 생성 (build 302).

    예: [["문화도시", "5급\\n6급\\n7급", "1\\n1\\n2", "업무1\\n...\\n업무6"]]
    → 3행 분리: "문화도시" rowspan=3, 직급/인원 각 줄 별도 행, 주요업무 균등 분배.

    판정 기준: 줄 수 > 1인 셀들의 최빈값(mode) = N → N행으로 분리.
    """
    from collections import Counter
    expanded = []

    for row in data:
        cells = [str(c or "") for c in row]
        line_lists = [c.split("\n") for c in cells]
        line_counts = [len(l) for l in line_lists]
        max_n = max(line_counts) if line_counts else 1

        if max_n <= 1:
            expanded.append([{"text": c, "rowspan": 1} for c in cells])
            continue

        multi = [c for c in line_counts if c > 1]
        if not multi:
            expanded.append([{"text": c, "rowspan": 1} for c in cells])
            continue

        target_n, freq = Counter(multi).most_common(1)[0]
        # 행 분리 조건 (build 305):
        # 1. 최소 2개 셀이 같은 줄 수 (freq >= 2)
        # 2. 1줄 셀 수가 동일줄 셀 수보다 적어야 (대부분이 줄바꿈인 경우만)
        #    예: "문화도시"(1줄) + 직급(3줄) + 인원(3줄) + 업무(6줄) → 1줄=1 < freq=2 → 분리 OK
        #    예: ""(1줄) + "7급\n보훈"(2줄) + "1"(1줄) + "업무"(2줄) → 1줄=2 >= freq=2 → 분리 안 함
        single_count = sum(1 for c in line_counts if c <= 1)
        if target_n <= 1 or freq < 2 or single_count >= freq:
            expanded.append([{"text": c, "rowspan": 1} for c in cells])
            continue

        for sub in range(target_n):
            new_row = []
            for col_idx, lines in enumerate(line_lists):
                n = len(lines)
                if n == target_n:
                    # 줄 수가 target과 정확히 같으면 → 각 줄이 별도 행
                    new_row.append({"text": lines[sub].strip(), "rowspan": 1})
                elif n == 1:
                    # 1줄 셀 → 첫 행에 rowspan=target_n
                    if sub == 0:
                        new_row.append({"text": lines[0].strip(), "rowspan": target_n})
                    else:
                        new_row.append(None)  # skip (rowspan으로 커버)
                else:
                    # 줄 수 불일치 → 하나의 셀에 rowspan으로 합침 (build 303)
                    # 원본 PDF에서 주요업무 등은 하나의 큰 셀로 병합되어 있음
                    if sub == 0:
                        new_row.append({"text": "\n".join(l.strip() for l in lines), "rowspan": target_n})
                    else:
                        new_row.append(None)  # skip (rowspan으로 커버)
            expanded.append(new_row)

    return expanded


def _merge_vertical_empty_cells(expanded):
    """수직으로 빈 셀이 연속되면 위쪽 셀에 rowspan 추가 (build 302).

    pdfplumber가 병합 셀의 하위 행을 빈 문자열로 반환하는 경우 처리.
    """
    if not expanded:
        return expanded

    num_cols = max(len(r) for r in expanded)
    result = [list(r) for r in expanded]  # deep copy

    # 첫 열(col=0)에만 적용 (build 306) — 채용공고에서 수직 병합은 카테고리(첫 열)에서만 발생.
    # 다른 열의 빈 셀은 "데이터 없음"이지 병합이 아님 (예: 발표일 빈 칸).
    for col in [0]:
        r = 0
        while r < len(result):
            cell = result[r][col] if col < len(result[r]) else None
            if cell is None or (isinstance(cell, dict) and not cell.get("text", "").strip()):
                r += 1
                continue

            # 아래로 빈 셀 연속 수 세기
            span = 1
            while r + span < len(result):
                below = result[r + span][col] if col < len(result[r + span]) else None
                if below is None:
                    span += 1
                elif isinstance(below, dict) and not below.get("text", "").strip():
                    result[r + span][col] = None  # skip으로 변환
                    span += 1
                else:
                    break

            if span > 1 and isinstance(cell, dict):
                cell["rowspan"] = span

            r += span

    return result


def _expanded_to_html(expanded, first_row_th=True):
    """확장된 데이터 → HTML 테이블 생성 (build 302)."""
    if not expanded:
        return ""
    parts = ["<table>"]
    for r_idx, row in enumerate(expanded):
        parts.append("  <tr>")
        for cell in row:
            if cell is None:
                continue
            tag = "th" if first_row_th and r_idx == 0 else "td"
            raw = cell.get("text", "") if isinstance(cell, dict) else str(cell)
            text = html_escape(raw).replace("\n", "<br>")  # 셀 내 줄바꿈 → <br> (build 304)
            rs = cell.get("rowspan", 1) if isinstance(cell, dict) else 1
            attrs = f' rowspan="{rs}"' if rs > 1 else ""
            parts.append(f"    <{tag}{attrs}>{text}</{tag}>")
        parts.append("  </tr>")
    parts.append("</table>")
    return "\n".join(parts)


def _build_table_html_with_spans(table_obj, page=None):
    """pdfplumber Table + Page에서 정확한 rowspan/colspan HTML 생성 (build 307).

    핵심: page.crop(cell_bbox).extract_text()로 각 셀 텍스트를 물리적 영역에서
    직접 추출. extract()의 2D 그리드 불일치를 완전히 우회.
    """
    try:
        cells = table_obj.cells  # [(x0, top, x1, bottom), ...]

        if not cells:
            return _simple_table_html(table_obj.extract())

        # 행/열 경계 추출 (tolerance 2px로 중복 제거)
        TOL = 2
        def _dedup(vals):
            vals = sorted(set(vals))
            result = [vals[0]] if vals else []
            for v in vals[1:]:
                if v - result[-1] > TOL:
                    result.append(v)
            return result

        y_coords = _dedup([c[1] for c in cells] + [c[3] for c in cells])
        x_coords = _dedup([c[0] for c in cells] + [c[2] for c in cells])

        def _find_idx(val, coords):
            for i, c in enumerate(coords):
                if abs(val - c) <= TOL:
                    return i
            return -1

        num_rows = len(y_coords) - 1
        num_cols = len(x_coords) - 1

        if num_rows <= 0 or num_cols <= 0:
            return _simple_table_html(data)

        # 셀별 그리드 위치 + span 계산
        grid = {}
        occupied = set()

        for cell_bbox in cells:
            x0, top, x1, bottom = cell_bbox
            r_start = _find_idx(top, y_coords)
            r_end = _find_idx(bottom, y_coords)
            c_start = _find_idx(x0, x_coords)
            c_end = _find_idx(x1, x_coords)

            if r_start < 0 or c_start < 0 or r_end < 0 or c_end < 0:
                continue

            rowspan = max(1, r_end - r_start)
            colspan = max(1, c_end - c_start)

            # 텍스트: data 2D 배열에서 가져오기 (병합 셀은 첫 위치에 텍스트)
            text = ""
            if r_start < len(data) and c_start < len(data[r_start]):
                text = str(data[r_start][c_start] or "")

            grid[(r_start, c_start)] = {
                "text": text,
                "rowspan": rowspan,
                "colspan": colspan,
            }

            for r in range(r_start, r_start + rowspan):
                for c in range(c_start, c_start + colspan):
                    if (r, c) != (r_start, c_start):
                        occupied.add((r, c))

        # HTML 생성
        parts = ["<table>"]
        for r in range(num_rows):
            parts.append("  <tr>")
            for c in range(num_cols):
                if (r, c) in occupied:
                    continue
                cell_info = grid.get((r, c))
                if not cell_info:
                    parts.append("    <td></td>")
                    continue

                tag = "th" if r == 0 else "td"
                val = html_escape(cell_info["text"])
                attrs = ""
                if cell_info["rowspan"] > 1:
                    attrs += f' rowspan="{cell_info["rowspan"]}"'
                if cell_info["colspan"] > 1:
                    attrs += f' colspan="{cell_info["colspan"]}"'
                parts.append(f"    <{tag}{attrs}>{val}</{tag}>")

            parts.append("  </tr>")
        parts.append("</table>")
        return "\n".join(parts)

    except Exception as e:
        # span 계산 실패 시 fallback
        return _simple_table_html(table_obj.extract() if hasattr(table_obj, 'extract') else [])


def _simple_table_html(data):
    """단순 테이블 HTML (rowspan/colspan 없음) — fallback용"""
    if not data:
        return ""
    parts = ["<table>"]
    for i, row in enumerate(data):
        parts.append("  <tr>")
        for cell in (row or []):
            tag = "th" if i == 0 else "td"
            val = html_escape(str(cell or ""))
            parts.append(f"    <{tag}>{val}</{tag}>")
        parts.append("  </tr>")
    parts.append("</table>")
    return "\n".join(parts)


def convert_pdf(data, filename=""):
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber가 설치되지 않았습니다. pip3 install pdfplumber")

    warnings = []
    html_parts = []
    text_parts = []

    with pdfplumber.open(BytesIO(data)) as pdf:
        total_pages = len(pdf.pages)
        if total_pages > 50:
            warnings.append(
                f"PDF가 {total_pages}페이지입니다. 처음 50페이지만 처리합니다."
            )

        for page_num, page in enumerate(pdf.pages[:50]):
            # 1. 테이블 추출 — find_tables()로 셀 bbox 기반 rowspan/colspan 지원 (build 301)
            table_bboxes = []
            found_tables = []
            try:
                found_tables = page.find_tables()
                table_bboxes = [ft.bbox for ft in found_tables]
            except Exception:
                pass

            for ft in found_tables:
                try:
                    data = ft.extract()
                    if data:
                        # 줄바꿈 → 행 분리 (문화도시: 5급/6급/7급 분리)
                        expanded = _expand_multiline_rows(data)
                        # 수직 빈 셀 → rowspan 병합 (관 광: 2행 병합)
                        merged = _merge_vertical_empty_cells(expanded)
                        html_parts.append(_expanded_to_html(merged))
                    else:
                        html_parts.append(_build_table_html_with_spans(ft))
                except Exception:
                    html_parts.append(_build_table_html_with_spans(ft))

            # 2. 텍스트 추출 (테이블 영역 제외)
            if table_bboxes:

                def not_within_tables(obj):
                    if "top" not in obj or "x0" not in obj:
                        return True
                    for bbox in table_bboxes:
                        tx0, ttop, tx1, tbottom = bbox
                        if (
                            obj["x0"] >= tx0 - 2
                            and obj["x0"] <= tx1 + 2
                            and obj["top"] >= ttop - 2
                            and obj["top"] <= tbottom + 2
                        ):
                            return False
                    return True

                try:
                    filtered_page = page.filter(not_within_tables)
                    text = filtered_page.extract_text()
                except Exception:
                    text = page.extract_text()  # filter 실패 시 전체 텍스트 폴백
            else:
                text = page.extract_text()

            if text:
                text_parts.append(text)
                for line in text.split("\n"):
                    line = line.strip()
                    if not line:
                        continue
                    if _is_heading_candidate(line):
                        html_parts.append(f"<h3>{html_escape(line)}</h3>")
                    else:
                        html_parts.append(f"<p>{html_escape(line)}</p>")

    html = "\n".join(html_parts)
    text = "\n".join(text_parts)

    return {
        "html": html,
        "text": text,
        "metadata": {
            "pages": total_pages,
            "warnings": warnings,
            "conversion_path": "pdf → pdfplumber → html",
        },
    }


# ============================================
# LibreOffice Path Detection
# ============================================


def find_libreoffice():
    """LibreOffice soffice 실행 경로 탐색"""
    candidates = [
        # macOS
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        # Linux
        "/usr/bin/soffice",
        "/usr/bin/libreoffice",
        "/usr/local/bin/soffice",
        # Windows
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    for path in candidates:
        if os.path.isfile(path):
            return path

    # PATH에서 탐색
    import shutil as _shutil

    found = _shutil.which("soffice") or _shutil.which("libreoffice")
    return found


SOFFICE_PATH = find_libreoffice()

# ============================================
# Vision 변환 프롬프트 상수
# ============================================

_VISION_SYSTEM_PROMPT = """당신은 한국 채용공고 문서를 HTML로 정확히 변환하는 전문가입니다.
규칙:
1. 원본의 모든 내용을 빠짐없이 변환 (생략 금지, 환각 금지)
2. 표의 셀 병합(rowspan, colspan)을 원본과 동일하게 재현
3. 번호 체계(1., ①, ◦, □, ※, -) 그대로 유지
4. 볼드는 <strong>, 밑줄은 <u>, 기울임은 <em>
5. <html><head><body><style> 태그 작성 금지 — 본문 HTML만 출력
6. 설명이나 주석 없이 HTML만 반환"""

_VISION_USER_PROMPT = "첨부된 채용공고 문서 이미지를 HTML로 변환해주세요. 모든 페이지의 모든 내용을 포함하고, 표의 셀 병합을 정확히 재현하세요."

# ============================================
# HWP → DOCX → HTML (LibreOffice 파이프라인)
# ============================================


def convert_any_to_pdf(data, ext, filename=""):
    """HWP/HWPX/DOCX/DOC → PDF 변환 (LibreOffice 사용). PDF bytes 반환."""
    if not SOFFICE_PATH:
        return None, "LibreOffice가 설치되지 않았습니다."

    import shutil

    tmp_dir = tempfile.mkdtemp()
    try:
        suffix = f".{ext}"
        in_path = os.path.join(tmp_dir, f"input{suffix}")
        with open(in_path, "wb") as f:
            f.write(data)

        result = subprocess.run(
            [
                SOFFICE_PATH,
                "--headless",
                "--convert-to",
                "pdf",
                in_path,
                "--outdir",
                tmp_dir,
            ],
            capture_output=True,
            text=True,
            timeout=90,
        )

        pdf_path = None
        if result.returncode == 0:
            real_tmp = os.path.realpath(tmp_dir)
            for search_dir in [tmp_dir, real_tmp]:
                candidate = os.path.join(search_dir, "input.pdf")
                if os.path.exists(candidate):
                    pdf_path = candidate
                    break
                pdf_files = [f for f in os.listdir(search_dir) if f.endswith(".pdf")]
                if pdf_files:
                    pdf_path = os.path.join(search_dir, pdf_files[0])
                    break

        if pdf_path:
            with open(pdf_path, "rb") as f:
                return f.read(), None

        stderr_info = result.stderr.strip() or result.stdout.strip() or "출력 없음"
        return None, f"PDF 변환 실패: {stderr_info}"

    except subprocess.TimeoutExpired:
        return None, "LibreOffice 변환 시간 초과 (90초)"
    except Exception as e:
        return None, f"PDF 변환 오류: {e}"
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


def convert_to_raw_html(data, ext):
    """HWP/HWPX/DOCX → LibreOffice HTML 직접 변환 (AI 처리 전 원문 HTML 추출용).
    Returns: (html_string, warnings_list)
    """
    import shutil

    warnings = []

    if not SOFFICE_PATH:
        return None, ["LibreOffice가 설치되지 않았습니다."]

    tmp_dir = tempfile.mkdtemp()
    try:
        suffix = f".{ext}"
        in_path = os.path.join(tmp_dir, f"input{suffix}")
        with open(in_path, "wb") as f:
            f.write(data)

        # LibreOffice → HTML 직접 변환
        result = subprocess.run(
            [
                SOFFICE_PATH,
                "--headless",
                "--convert-to",
                "html",
                in_path,
                "--outdir",
                tmp_dir,
            ],
            capture_output=True,
            timeout=90,
        )

        html_path = None
        if result.returncode == 0:
            real_tmp = os.path.realpath(tmp_dir)
            for search_dir in [tmp_dir, real_tmp]:
                # input.html 먼저 탐색
                candidate = os.path.join(search_dir, "input.html")
                if os.path.exists(candidate):
                    html_path = candidate
                    break
                html_files = [f for f in os.listdir(search_dir) if f.endswith(".html")]
                if html_files:
                    html_path = os.path.join(search_dir, html_files[0])
                    break

        if not html_path:
            stderr_info = (
                result.stderr.decode("utf-8", errors="replace").strip() or "출력 없음"
            )
            return None, [f"LibreOffice HTML 변환 실패: {stderr_info}"]

        # 인코딩 자동 감지 (EUC-KR / UTF-8)
        raw_bytes = open(html_path, "rb").read()
        for enc in ("utf-8", "cp949", "euc-kr", "utf-8-sig"):
            try:
                html_content = raw_bytes.decode(enc)
                break
            except Exception:
                continue
        else:
            html_content = raw_bytes.decode("utf-8", errors="replace")

        # <body> 내용만 추출
        body_match = re.search(
            r"<body[^>]*>(.*?)</body>", html_content, re.DOTALL | re.IGNORECASE
        )
        body_html = body_match.group(1).strip() if body_match else html_content

        # LibreOffice 생성 노이즈 정리
        # 1. 페이지 구분선 제거
        body_html = re.sub(r"<hr\s*/?>", "", body_html, flags=re.IGNORECASE)
        # 2. 빈 단락 연속 축소 (3개 이상 → 2개)
        body_html = re.sub(
            r"(<p[^>]*>\s*<br\s*/?>\s*</p>\s*){3,}", "<p><br></p><p><br></p>", body_html
        )
        # 3. LibreOffice 전용 class 정리 (내용 유지)
        body_html = re.sub(
            r'\s+class="(P[0-9]+|T[0-9]+|Table[0-9]+|Frame[0-9]+)"', "", body_html
        )
        # 4. 불필요한 font-size:0 스팬 제거
        body_html = re.sub(r"<span[^>]*font-size:\s*0[^>]*>\s*</span>", "", body_html)

        warnings.append(f"LibreOffice HTML 변환 완료 ({ext.upper()} → HTML)")
        return body_html, warnings

    except subprocess.TimeoutExpired:
        return None, ["LibreOffice 변환 시간 초과 (90초)"]
    except Exception as e:
        return None, [f"원문 HTML 변환 오류: {e}"]
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


def convert_hwp_via_libreoffice(data):
    """LibreOffice로 HWP → DOCX 변환 후 python-docx로 HTML 추출, 실패 시 HTML 직접 변환"""
    if not SOFFICE_PATH:
        return None, "LibreOffice가 설치되지 않았습니다."

    import shutil

    tmp_dir = tempfile.mkdtemp()

    try:
        # 1. HWP 임시 파일 저장
        hwp_path = os.path.join(tmp_dir, "input.hwp")
        with open(hwp_path, "wb") as f:
            f.write(data)

        # 2A. 1차 시도: HWP → DOCX (테이블 보존 최적)
        result = subprocess.run(
            [
                SOFFICE_PATH,
                "--headless",
                "--convert-to",
                "docx:MS Word 2007 XML",
                hwp_path,
                "--outdir",
                tmp_dir,
            ],
            capture_output=True,
            text=True,
            timeout=60,
        )

        docx_path = None
        if result.returncode == 0:
            # realpath로 심링크 해소 (macOS /var → /private/var)
            real_tmp = os.path.realpath(tmp_dir)
            for search_dir in [tmp_dir, real_tmp]:
                candidate = os.path.join(search_dir, "input.docx")
                if os.path.exists(candidate):
                    docx_path = candidate
                    break
                docx_files = [f for f in os.listdir(search_dir) if f.endswith(".docx")]
                if docx_files:
                    docx_path = os.path.join(search_dir, docx_files[0])
                    break

        if docx_path:
            # python-docx로 HTML 추출 (테이블 보존)
            with open(docx_path, "rb") as f:
                docx_data = f.read()
            result = convert_docx(docx_data, "converted.docx")
            result["metadata"]["warnings"] = result["metadata"].get("warnings", [])
            result["metadata"]["warnings"].insert(
                0, "HWP → DOCX(LibreOffice) → HTML 변환 완료"
            )
            result["metadata"][
                "conversion_path"
            ] = "hwp → libreoffice → docx → python-docx → html"
            return result, None

        # 2B. 2차 시도: HWP → HTML 직접 변환 (DOCX 실패 시 fallback)
        print(
            f"[HWP] DOCX 변환 실패, HTML 직접 변환 시도 (stdout: {result.stdout.strip()}, stderr: {result.stderr.strip()})"
        )
        result_html = subprocess.run(
            [
                SOFFICE_PATH,
                "--headless",
                "--convert-to",
                "html",
                hwp_path,
                "--outdir",
                tmp_dir,
            ],
            capture_output=True,
            text=True,
            timeout=60,
        )

        html_path = None
        if result_html.returncode == 0:
            real_tmp = os.path.realpath(tmp_dir)
            for search_dir in [tmp_dir, real_tmp]:
                html_files = [f for f in os.listdir(search_dir) if f.endswith(".html")]
                if html_files:
                    html_path = os.path.join(search_dir, html_files[0])
                    break

        if html_path:
            with open(html_path, "r", encoding="utf-8", errors="replace") as f:
                html_content = f.read()
            # <body> 내부만 추출
            body_match = re.search(
                r"<body[^>]*>(.*?)</body>", html_content, re.DOTALL | re.IGNORECASE
            )
            html = body_match.group(1).strip() if body_match else html_content
            text = re.sub(r"<[^>]+>", " ", html)
            text = re.sub(r"\s+", " ", text).strip()
            return {
                "html": html,
                "text": text,
                "metadata": {
                    "warnings": [
                        "HWP → HTML(LibreOffice) 직접 변환 완료 (DOCX 경로 실패)"
                    ],
                    "conversion_path": "hwp → libreoffice → html",
                },
            }, None

        # 두 경로 모두 실패
        stderr_info = result.stderr.strip() or result_html.stderr.strip() or "출력 없음"
        return None, f"LibreOffice HWP 변환 실패 (DOCX·HTML 모두): {stderr_info}"

    except subprocess.TimeoutExpired:
        return None, "LibreOffice 변환 시간 초과 (60초)"
    except Exception as e:
        return None, f"LibreOffice 파이프라인 오류: {e}"
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


# ============================================
# HWP Vision 변환 (HWP → PDF → 이미지 → Claude Vision → HTML)
# ============================================


def convert_hwp_vision(data: bytes, filename: str = "", api_key: str = ""):
    """HWP → PDF → 이미지 → Claude Vision → HTML.
    api_key 없거나 변환 실패 시 None 반환 → 호출부에서 fallback 처리.
    """
    import anthropic
    import base64
    import io

    if not api_key:
        return None

    if not SOFFICE_PATH:
        return None

    try:
        from pdf2image import convert_from_path
        from PIL import Image
    except ImportError as e:
        print(f"[Vision] 의존성 없음 (pdf2image/Pillow): {e}")
        return None

    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            safe_filename = os.path.basename(filename) if filename else "input.hwp"
            hwp_path = os.path.join(tmpdir, safe_filename or "input.hwp")
            with open(hwp_path, "wb") as f:
                f.write(data)

            # 1. HWP → PDF (LibreOffice)
            result = subprocess.run(
                [SOFFICE_PATH, "--headless", "--convert-to", "pdf", "--outdir", tmpdir, hwp_path],
                capture_output=True,
                timeout=120,
            )
            if result.returncode != 0:
                print(f"[Vision] LibreOffice PDF 변환 실패: {result.stderr.decode(errors='replace')}")
                return None

            from pathlib import Path as _Path
            pdf_path = os.path.join(tmpdir, _Path(hwp_path).stem + ".pdf")
            if not os.path.exists(pdf_path):
                print("[Vision] PDF 파일 없음")
                return None

            # 2. PDF → 이미지
            pages = convert_from_path(pdf_path, dpi=300)
            images_b64 = []
            for page in pages:
                max_dim = 2048
                if max(page.size) > max_dim:
                    ratio = max_dim / max(page.size)
                    page = page.resize(
                        (int(page.size[0] * ratio), int(page.size[1] * ratio)),
                        Image.LANCZOS,
                    )
                buf = io.BytesIO()
                page.save(buf, format="PNG", optimize=True)
                images_b64.append(base64.standard_b64encode(buf.getvalue()).decode())

            # 3. Claude Vision API 호출 (tmpdir 블록 안에서 완료)
            client = anthropic.Anthropic(api_key=api_key)
            content = []
            for img_b64 in images_b64:
                content.append({
                    "type": "image",
                    "source": {"type": "base64", "media_type": "image/png", "data": img_b64},
                })
            content.append({"type": "text", "text": _VISION_USER_PROMPT})

            response = client.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=16000,
                temperature=0,
                system=_VISION_SYSTEM_PROMPT,
                messages=[{"role": "user", "content": content}],
            )

            html = response.content[0].text.strip()
            text = re.sub(r"<[^>]+>", "", html)
            text = re.sub(r"\s+", " ", text).strip()

            return {
                "html": html,
                "text": text,
                "metadata": {
                    "format": "hwp-vision",
                    "pages": len(pages),
                    "warnings": ["Claude Vision으로 변환 완료"],
                },
            }

    except Exception as e:
        print(f"[Vision] 변환 예외: {e}")
        return None


# ============================================
# HWP Converter (LibreOffice → hwp5 → textutil → fallback)
# ============================================


def convert_hwp(data, filename="", parser_hint="", api_key=""):
    warnings = []

    # Vision 최우선: api_key 있고 parser_hint != "native"
    # api_key 없으면 환경변수 fallback
    effective_key = api_key or os.environ.get("ANTHROPIC_API_KEY", "")
    if parser_hint != "native" and effective_key:
        vision_result = convert_hwp_vision(data, filename, api_key=effective_key)
        if vision_result:
            return vision_result

    if parser_hint == "native":
        # v2 경로: 네이티브 파서 1순위
        try:
            native_result = parse_hwp5_native(data)
            if native_result and native_result.get("text", "").strip():
                native_warnings = native_result.get("metadata", {}).get("warnings", [])
                native_warnings.insert(0, "HWP 5.0 네이티브 파서로 변환 완료 (v2 1순위)")
                native_result["metadata"]["warnings"] = warnings + native_warnings
                return native_result
            if native_result:
                warnings.append("네이티브 파서: 텍스트 추출 결과가 비어있음")
        except Exception as e:
            warnings.append(f"네이티브 HWP 파서 실패: {e}")
            print(f"[HWP] 네이티브 파서 예외 (v2): {e}")

        # 네이티브 실패 → LibreOffice fallback
        if SOFFICE_PATH:
            lo_result, lo_error = convert_hwp_via_libreoffice(data)
            if lo_result:
                lo_result.setdefault("metadata", {})
                lo_result["metadata"].setdefault("warnings", [])
                lo_result["metadata"]["warnings"] = warnings + lo_result["metadata"]["warnings"]
                return lo_result
            if lo_error:
                warnings.append(lo_error)
    else:
        # 기본 경로: LibreOffice 1순위 (기존 동작 100% 유지)
        # build 309: 원래 순서 복구. 네이티브 파서가 이 HWP에서 표를 못 잡고
        # LibreOffice도 컨테이너 재시작 후 실패하는 문제 발견 → 안정성 우선.
        if SOFFICE_PATH:
            lo_result, lo_error = convert_hwp_via_libreoffice(data)
            if lo_result:
                return lo_result
            if lo_error:
                warnings.append(lo_error)

        try:
            native_result = parse_hwp5_native(data)
            if native_result and native_result.get("text", "").strip():
                native_warnings = native_result.get("metadata", {}).get("warnings", [])
                native_warnings.insert(0, "HWP 5.0 네이티브 파서로 변환 완료")
                native_result["metadata"]["warnings"] = warnings + native_warnings
                return native_result
            if native_result:
                warnings.append("네이티브 파서: 텍스트 추출 결과가 비어있음")
        except Exception as e:
            warnings.append(f"네이티브 HWP 파서 실패: {e}")
            print(f"[HWP] 네이티브 파서 예외: {e}")

    # 3순위: 바이너리 텍스트 추출 폴백 (양쪽 경로 공통)
    warnings.append(
        "HWP 파일은 제한적으로 지원됩니다. DOCX로 변환 후 업로드를 권장합니다."
    )
    text = extract_text_from_binary(data)
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    html = "\n".join(f"<p>{html_escape(l)}</p>" for l in lines)
    return {"html": html, "text": "\n".join(lines), "metadata": {"warnings": warnings}}


# ============================================
# HWPX Converter (ZIP-based)
# ============================================


def convert_hwpx(data, filename=""):
    warnings = []

    try:
        with zipfile.ZipFile(BytesIO(data)) as zf:
            # 모든 section XML 수집 (순서대로)
            section_xmls = []
            for name in sorted(zf.namelist()):
                if "section" in name.lower() and name.endswith(".xml"):
                    section_xmls.append(zf.read(name).decode("utf-8", errors="replace"))

            if not section_xmls:
                for name in zf.namelist():
                    if name.endswith(".xml") and "meta" not in name.lower():
                        section_xmls.append(
                            zf.read(name).decode("utf-8", errors="replace")
                        )
                        break

            if section_xmls:
                html_parts = []
                text_parts = []
                for xml_content in section_xmls:
                    h, t = _parse_hwpx_section(xml_content, warnings)
                    if h:
                        html_parts.append(h)
                    if t:
                        text_parts.append(t)

                html = "\n".join(html_parts)
                text = "\n".join(text_parts)

                if html.strip():
                    return {
                        "html": html,
                        "text": text,
                        "metadata": {"warnings": warnings},
                    }

            # XML 파싱이 빈 결과 → regex 폴백
            warnings.append("HWPX XML 구조 파싱 실패. regex 폴백 사용.")
            return _convert_hwpx_regex_fallback(data, warnings)

    except zipfile.BadZipFile:
        warnings.append("유효하지 않은 HWPX(ZIP) 파일입니다.")

    # 폴백
    return _convert_hwpx_regex_fallback(data, warnings)


def _parse_hwpx_section(xml_content, warnings):
    """HWPX section XML을 파싱하여 (html, text) 반환."""
    import xml.etree.ElementTree as ET

    try:
        root = ET.fromstring(xml_content)
    except ET.ParseError as e:
        warnings.append(f"HWPX XML 파싱 오류: {e}")
        # regex 폴백
        texts = re.findall(r"<hp:t[^>]*>([^<]+)</hp:t>", xml_content)
        if texts:
            html = "\n".join(
                f"<p>{html_escape(t.strip())}</p>" for t in texts if t.strip()
            )
            text = "\n".join(t.strip() for t in texts if t.strip())
            return html, text
        return "", ""

    html_parts = []
    text_parts = []

    def _local(tag):
        """네임스페이스 제거하여 로컬 태그명 반환."""
        return tag.split("}")[-1] if "}" in tag else tag

    def _extract_text(elem):
        """요소 내부의 모든 't' 텍스트 수집."""
        texts = []
        for sub in elem.iter():
            if _local(sub.tag) == "t" and sub.text:
                texts.append(sub.text)
        return "".join(texts)

    def _extract_cell_paras(tc_elem):
        """셀 내부 직계 자식 단락만 반환 (중첩 테이블 내부 p 중복 방지)."""
        paras = []
        nested_html = []
        for sub in tc_elem:  # 직계 자식만 순회 (.iter() 대신)
            if _local(sub.tag) == "p":
                p_text = _extract_text(sub)
                paras.append(p_text)
            elif _local(sub.tag) == "tbl":
                tbl_h, _ = _parse_hwpx_table(
                    sub, _local, _extract_text, _extract_cell_paras
                )
                if tbl_h:
                    nested_html.append(tbl_h)
        # 후행 빈 단락 제거
        while paras and not paras[-1].strip():
            paras.pop()
        return paras, nested_html

    def _hwpx_classify_table(tbl_elem):
        """ki-it.com 논문: HWPX 테이블 lettered/unlettered 분류.
        - 단일 셀: unlettered
        - 1행 테이블: unlettered (HWP 텍스트박스·레이아웃 표)
        - 비어있는 셀 비율이 85% 초과: unlettered
        """
        total_cells = 0
        non_empty_cells = 0
        num_rows = 0
        for tr in tbl_elem:
            if _local(tr.tag) != "tr":
                continue
            num_rows += 1
            for tc in tr:
                if _local(tc.tag) != "tc":
                    continue
                total_cells += 1
                cell_text = _extract_text(tc).strip()
                if cell_text:
                    non_empty_cells += 1
        if total_cells <= 1:
            return "unlettered"
        # 1행 테이블: HWP 텍스트박스·2단 레이아웃 → 레이아웃 표
        if num_rows == 1:
            return "unlettered"
        return "lettered" if (non_empty_cells / total_cells) >= 0.15 else "unlettered"

    # 루트 직접 자식 순회 (문서 순서 보존)
    for child in root:
        local = _local(child.tag)

        if local == "tbl":
            tbl_type = _hwpx_classify_table(child)
            if tbl_type == "unlettered":
                # 레이아웃 표: 셀 텍스트를 단락으로 변환
                for tr in child:
                    if _local(tr.tag) != "tr":
                        continue
                    for tc in tr:
                        if _local(tc.tag) != "tc":
                            continue
                        paras, _ = _extract_cell_paras(tc)
                        cell_text = "\n".join(p for p in paras if p.strip())
                        if cell_text.strip():
                            html_parts.append(f"<p>{html_escape(cell_text)}</p>")
                            text_parts.append(cell_text)
            else:
                tbl_html, tbl_text = _parse_hwpx_table(
                    child, _local, _extract_text, _extract_cell_paras
                )
                html_parts.append(tbl_html)
                text_parts.append(tbl_text)

        elif local == "p":
            p_text = _extract_text(child)
            if p_text.strip():
                html_parts.append(f"<p>{html_escape(p_text)}</p>")
                text_parts.append(p_text)

    return "\n".join(html_parts), "\n".join(text_parts)


def _parse_hwpx_table(tbl_elem, _local, _extract_text, _extract_cell_paras):
    """HWPX hp:tbl 요소를 HTML 테이블로 변환 (정렬·여백·배경색 inline style 포함)."""
    html_parts = ["<table>"]
    text_parts = []
    row_idx = 0

    def _hwpx_tc_style(tc_elem):
        """HWPX tc 요소에서 셀 스타일 CSS 문자열 추출."""
        styles = []
        # tc 직접 속성 (일부 구현체)
        h_align = (
            tc_elem.get("hAlign") or tc_elem.get("align") or tc_elem.get("halign", "")
        )
        v_align = tc_elem.get("vAlign") or tc_elem.get("verticalAlign", "")
        h_map = {
            "left": "left",
            "center": "center",
            "right": "right",
            "justify": "justify",
            "distribute": "justify",
        }
        v_map = {
            "top": "top",
            "center": "middle",
            "middle": "middle",
            "bottom": "bottom",
        }
        if h_align in h_map:
            styles.append(f"text-align:{h_map[h_align]}")
        if v_align in v_map:
            styles.append(f"vertical-align:{v_map[v_align]}")

        for sub in tc_elem:
            local = _local(sub.tag)
            # cellPr / tcPr 자식 요소
            if local in ("cellPr", "tcPr"):
                h_val = sub.get("hAlign") or sub.get("align", "")
                v_val = sub.get("vAlign") or sub.get("verticalAlign", "")
                if h_val in h_map:
                    styles.append(f"text-align:{h_map[h_val]}")
                if v_val in v_map:
                    styles.append(f"vertical-align:{v_map[v_val]}")
                # 배경색
                bg = sub.get("fillColor") or sub.get("bgColor") or sub.get("color", "")
                if bg and bg not in ("FFFFFF", "ffffff", "none", ""):
                    if not bg.startswith("#"):
                        bg = f"#{bg}"
                    styles.append(f"background-color:{bg}")
            # tcMargin / margin
            elif local in ("tcMargin", "margin", "cellMargin"):
                try:
                    left = sub.get("left", "")
                    right = sub.get("right", "")
                    top = sub.get("top", "")
                    bottom = sub.get("bottom", "")
                    pads = []
                    for v in [top, right, bottom, left]:
                        if v:
                            # HWPX 단위는 보통 1/100mm → px (1mm ≈ 3.78px)
                            pads.append(f"{int(v) * 3.78 / 100:.0f}px")
                        else:
                            pads.append("0")
                    if any(p != "0" for p in pads):
                        styles.append(f'padding:{" ".join(pads)}')
                except (ValueError, TypeError):
                    pass
            # tcBkg / background
            elif local in ("tcBkg", "background", "cellBkg"):
                bg = sub.get("color") or sub.get("fillColor", "")
                if bg and bg not in ("FFFFFF", "ffffff", "none", ""):
                    if not bg.startswith("#"):
                        bg = f"#{bg}"
                    styles.append(f"background-color:{bg}")

        return ";".join(styles)

    for tr_elem in tbl_elem:
        if _local(tr_elem.tag) != "tr":
            continue

        html_parts.append("  <tr>")
        for tc_elem in tr_elem:
            if _local(tc_elem.tag) != "tc":
                continue

            # cellSpan에서 colspan/rowspan 추출
            colspan = 1
            rowspan = 1
            for sub in tc_elem:
                if _local(sub.tag) == "cellSpan":
                    try:
                        colspan = int(sub.get("colSpan", "1"))
                        rowspan = int(sub.get("rowSpan", "1"))
                    except (ValueError, TypeError):
                        pass

            paras, nested_html = _extract_cell_paras(tc_elem)
            cell_text = "\n".join(paras)
            text_parts.append(cell_text)

            cell_tag = "th" if row_idx == 0 else "td"
            cell_html = html_escape(cell_text).replace("\n", "<br>")
            if nested_html:
                cell_html += "".join(nested_html)
            attrs = ""
            if colspan > 1:
                attrs += f' colspan="{colspan}"'
            if rowspan > 1:
                attrs += f' rowspan="{rowspan}"'
            cell_style = _hwpx_tc_style(tc_elem)
            if cell_style:
                attrs += f' style="{cell_style}"'
            html_parts.append(f"    <{cell_tag}{attrs}>{cell_html}</{cell_tag}>")

        html_parts.append("  </tr>")
        row_idx += 1

    html_parts.append("</table>")
    return "\n".join(html_parts), "\n".join(text_parts)


def _convert_hwpx_regex_fallback(data, warnings):
    """기존 regex 기반 HWPX 텍스트 추출 (폴백)."""
    try:
        with zipfile.ZipFile(BytesIO(data)) as zf:
            for name in zf.namelist():
                if "section" in name.lower() and name.endswith(".xml"):
                    xml_content = zf.read(name).decode("utf-8", errors="replace")
                    texts = re.findall(r"<hp:t[^>]*>([^<]+)</hp:t>", xml_content)
                    if not texts:
                        texts = re.findall(
                            r"<(?:t|text)[^>]*>([^<]+)</(?:t|text)>", xml_content
                        )
                    if texts:
                        html = "\n".join(
                            f"<p>{html_escape(t.strip())}</p>"
                            for t in texts
                            if t.strip()
                        )
                        text = "\n".join(t.strip() for t in texts if t.strip())
                        return {
                            "html": html,
                            "text": text,
                            "metadata": {"warnings": warnings},
                        }
    except Exception:
        pass

    text = extract_text_from_binary(data)
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    html = "\n".join(f"<p>{html_escape(l)}</p>" for l in lines)
    warnings.append("HWPX 변환이 제한적입니다. DOCX로 변환 후 다시 시도해주세요.")
    return {"html": html, "text": "\n".join(lines), "metadata": {"warnings": warnings}}


# ============================================
# Text/CSV/HTML Converters
# ============================================


def convert_text(data, filename=""):
    text = data.decode("utf-8", errors="replace")
    lines = text.split("\n")
    html = "\n".join(f"<p>{html_escape(l)}</p>" if l.strip() else "<br>" for l in lines)
    return {"html": html, "text": text, "metadata": {"warnings": []}}


def convert_csv(data, filename=""):
    import csv

    text = data.decode("utf-8", errors="replace")
    reader = csv.reader(text.splitlines())
    rows = list(reader)

    html_parts = ["<table>"]
    for i, row in enumerate(rows):
        html_parts.append("  <tr>")
        for cell in row:
            tag = "th" if i == 0 else "td"
            html_parts.append(f"    <{tag}>{html_escape(cell)}</{tag}>")
        html_parts.append("  </tr>")
    html_parts.append("</table>")

    html = "\n".join(html_parts)
    return {"html": html, "text": text, "metadata": {"rows": len(rows), "warnings": []}}


def convert_html_file(data, filename=""):
    html = data.decode("utf-8", errors="replace")
    body_match = re.search(r"<body[^>]*>(.*?)</body>", html, re.DOTALL | re.IGNORECASE)
    if body_match:
        html = body_match.group(1)
    text = re.sub(r"<[^>]+>", " ", html)
    text = re.sub(r"\s+", " ", text).strip()
    return {"html": html, "text": text, "metadata": {"warnings": []}}


# ============================================
# Utility
# ============================================


def extract_text_from_binary(data):
    """바이너리 데이터에서 인쇄 가능한 텍스트 추출"""
    # UTF-16LE 시도 (HWP)
    try:
        text = data.decode("utf-16-le", errors="replace")
    except:
        text = data.decode("utf-8", errors="replace")

    # 인쇄 가능한 문자만 보존
    text = re.sub(
        r"[^\x20-\x7E\uAC00-\uD7AF\u3131-\u318E\n\r\t.,;:!?\-()[\]{}\'\"\/\\@#$%&*+=<>~`]",
        " ",
        text,
    )
    text = re.sub(r"\s{3,}", "\n", text)
    return text.strip()


# ============================================
# HWP 5.0 Native Parser (stdlib only: struct + zlib)
# OLE2 Compound Binary → HWP streams → text/table extraction
# ============================================

OLE2_SIGNATURE = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"
ENDOFCHAIN = 0xFFFFFFFE
FREESECT = 0xFFFFFFFF
NOSTREAM = 0xFFFFFFFF

# HWP record tag IDs
HWPTAG_PARA_HEADER = 66
HWPTAG_PARA_TEXT = 67
HWPTAG_PARA_CHAR_SHAPE = 68   # 단락 내 문자 모양 ID 매핑
HWPTAG_CTRL_HEADER = 71
HWPTAG_LIST_HEADER = 72
HWPTAG_TABLE = 77
HWPTAG_DOCINFO_CHAR_SHAPE = 21  # DocInfo 스트림 내 CharShape (0x15)
HWPTAG_DOCINFO_PARA_SHAPE = 25  # DocInfo 스트림 내 ParaShape (0x19)

# Control chars that consume 16 bytes in PARA_TEXT
_CTRL_16BYTE = frozenset(
    [1, 2, 3, 4, 5, 6, 7, 8, 9, 11, 12, 14, 15, 16, 17, 18, 19, 20, 21, 22, 23]
)


class OLE2Reader:
    """Minimal OLE2 Compound Binary File reader (Python stdlib only)."""

    def __init__(self, data):
        self.data = data
        if data[:8] != OLE2_SIGNATURE:
            raise ValueError("OLE2 파일이 아닙니다.")

        # Parse header
        self.sector_shift = struct.unpack_from("<H", data, 30)[0]
        self.sector_size = 1 << self.sector_shift
        self.mini_sector_shift = struct.unpack_from("<H", data, 32)[0]
        self.mini_sector_size = 1 << self.mini_sector_shift

        self.num_fat_sectors = struct.unpack_from("<I", data, 44)[0]
        self.first_dir_sid = struct.unpack_from("<I", data, 48)[0]
        self.mini_stream_cutoff = struct.unpack_from("<I", data, 56)[0]
        self.first_minifat_sid = struct.unpack_from("<I", data, 60)[0]
        self.num_minifat_sectors = struct.unpack_from("<I", data, 64)[0]
        self.first_difat_sid = struct.unpack_from("<I", data, 68)[0]
        self.num_difat_sectors = struct.unpack_from("<I", data, 72)[0]

        # Build FAT
        self.fat = self._build_fat()
        # Parse directory
        self.entries = self._read_directory()
        # Build mini FAT and mini stream (lazy)
        self._mini_fat = None
        self._mini_stream = None

    def _sector_data(self, sid):
        offset = (sid + 1) * self.sector_size
        return self.data[offset : offset + self.sector_size]

    def _get_chain(self, start_sid):
        chain = []
        sid = start_sid
        seen = set()
        while sid < ENDOFCHAIN and sid not in seen:
            seen.add(sid)
            chain.append(sid)
            if sid >= len(self.fat):
                break
            sid = self.fat[sid]
        return chain

    def _build_fat(self):
        # DIFAT: first 109 entries from header
        difat = list(struct.unpack_from("<109I", self.data, 76))
        difat = [s for s in difat if s < ENDOFCHAIN]

        # Additional DIFAT sectors
        if self.num_difat_sectors > 0 and self.first_difat_sid < ENDOFCHAIN:
            sid = self.first_difat_sid
            ipp = self.sector_size // 4 - 1  # entries per sector (last = next ptr)
            seen = set()
            for _ in range(self.num_difat_sectors):
                if sid >= ENDOFCHAIN or sid in seen:
                    break
                seen.add(sid)
                sec = self._sector_data(sid)
                entries = struct.unpack_from(f"<{ipp}I", sec, 0)
                difat.extend(s for s in entries if s < ENDOFCHAIN)
                sid = struct.unpack_from("<I", sec, ipp * 4)[0]

        # Read FAT sectors
        fat = []
        epp = self.sector_size // 4
        for fat_sid in difat[: self.num_fat_sectors]:
            sec = self._sector_data(fat_sid)
            fat.extend(struct.unpack_from(f"<{epp}I", sec, 0))
        return fat

    def _read_directory(self):
        chain = self._get_chain(self.first_dir_sid)
        raw = b"".join(self._sector_data(sid) for sid in chain)

        entries = []
        for i in range(0, len(raw), 128):
            entry_data = raw[i : i + 128]
            if len(entry_data) < 128:
                break
            name_size = struct.unpack_from("<H", entry_data, 64)[0]
            if name_size <= 0 or name_size > 64:
                entries.append(None)
                continue
            try:
                name = entry_data[: name_size - 2].decode("utf-16le")
            except:
                name = ""
            obj_type = entry_data[66]
            child = struct.unpack_from("<I", entry_data, 76)[0]
            left = struct.unpack_from("<I", entry_data, 68)[0]
            right = struct.unpack_from("<I", entry_data, 72)[0]
            start_sid = struct.unpack_from("<I", entry_data, 116)[0]
            size = struct.unpack_from("<I", entry_data, 120)[0]
            entries.append(
                {
                    "name": name,
                    "type": obj_type,
                    "child": child,
                    "left": left,
                    "right": right,
                    "start_sid": start_sid,
                    "size": size,
                    "index": len(entries),
                }
            )
        return entries

    def _ensure_mini(self):
        if self._mini_fat is not None:
            return
        # Mini FAT
        self._mini_fat = []
        if self.first_minifat_sid < ENDOFCHAIN:
            chain = self._get_chain(self.first_minifat_sid)
            epp = self.sector_size // 4
            for sid in chain:
                sec = self._sector_data(sid)
                self._mini_fat.extend(struct.unpack_from(f"<{epp}I", sec, 0))
        # Mini stream = root entry's data chain
        root = self.entries[0] if self.entries else None
        if root and root["start_sid"] < ENDOFCHAIN:
            chain = self._get_chain(root["start_sid"])
            self._mini_stream = b"".join(self._sector_data(sid) for sid in chain)
        else:
            self._mini_stream = b""

    def _get_mini_chain(self, start_sid):
        chain = []
        sid = start_sid
        seen = set()
        while sid < ENDOFCHAIN and sid not in seen:
            seen.add(sid)
            chain.append(sid)
            if sid >= len(self._mini_fat):
                break
            sid = self._mini_fat[sid]
        return chain

    def read_stream(self, entry):
        """Read the data of a directory entry (stream)."""
        if entry is None or entry["type"] != 2:
            return b""
        size = entry["size"]
        if size == 0:
            return b""

        if size < self.mini_stream_cutoff:
            # Read from mini stream
            self._ensure_mini()
            chain = self._get_mini_chain(entry["start_sid"])
            parts = []
            for msid in chain:
                off = msid * self.mini_sector_size
                parts.append(self._mini_stream[off : off + self.mini_sector_size])
            return b"".join(parts)[:size]
        else:
            chain = self._get_chain(entry["start_sid"])
            return b"".join(self._sector_data(sid) for sid in chain)[:size]

    def _walk_tree(self, did, parent_path=""):
        """Walk the directory red-black tree (in-order traversal)."""
        if did == NOSTREAM or did >= len(self.entries):
            return
        entry = self.entries[did]
        if entry is None:
            return
        # Left
        yield from self._walk_tree(entry["left"], parent_path)
        # Self
        full_path = parent_path + "/" + entry["name"] if parent_path else entry["name"]
        yield full_path, entry
        # Children (sub-tree)
        if entry["type"] in (1, 5) and entry["child"] != NOSTREAM:
            yield from self._walk_tree(entry["child"], full_path)
        # Right
        yield from self._walk_tree(entry["right"], parent_path)

    def list_entries(self):
        """List all entries with full paths."""
        root = self.entries[0] if self.entries else None
        if not root:
            return []
        result = [("Root Entry", root)]
        if root["child"] != NOSTREAM:
            result.extend(self._walk_tree(root["child"], "Root Entry"))
        return result

    def find_entry(self, name_lower):
        """Find entry by name (case-insensitive)."""
        for path, entry in self.list_entries():
            if entry["name"].lower() == name_lower:
                return entry
        return None

    def find_entries_by_prefix(self, prefix_lower):
        """Find all entries whose path contains the prefix."""
        results = []
        for path, entry in self.list_entries():
            if prefix_lower in path.lower():
                results.append((path, entry))
        return results


def _parse_docinfo_charshapes(ole, compressed):
    """OLE2 DocInfo 스트림에서 CharShape 목록 파싱.
    반환: {shape_idx(int): {'bold': bool, 'underline': bool}} dict
    CharShape 구조 (HWP5):
      FaceNameId[7]×2=14 + Sz[7]×4=28 + CharSpacing[7]=7
      + RelSz[7]=7 + CharOffset[7]=7 + Ratio[7]=7 = 70bytes
      + attr UINT32 at offset 70 (bit1=bold, bit2=underline)
    """
    docinfo_entry = ole.find_entry("docinfo")
    if not docinfo_entry:
        return {}
    raw = ole.read_stream(docinfo_entry)
    if not raw:
        return {}
    if compressed:
        for wbits in (-15, 15, 15 + 32):
            try:
                raw = zlib.decompress(raw, wbits)
                break
            except zlib.error:
                continue
    records = _parse_records(raw)
    charshape_map = {}
    idx = 0
    for rec in records:
        if rec["tag"] == HWPTAG_DOCINFO_CHAR_SHAPE:
            pl = rec["payload"]
            bold = False
            underline = False
            # CharShape 레이아웃:
            #   [0-13]  FaceNameId[7]×WORD=14  [14-20] Ratio[7]×BYTE=7
            #   [21-27] CharSpacing[7]×BYTE=7  [28-34] RelSz[7]×BYTE=7
            #   [35-41] CharOffset[7]×BYTE=7   [42-45] BaseSz UINT32
            #   [46-49] CharShapeAttr UINT32 (bit1=bold, bit2=underline)
            #   [52-55] CharColor COLORREF (0x00BBGGRR little-endian)
            color = None
            if len(pl) >= 50:
                attr = struct.unpack_from("<I", pl, 46)[0]
                bold = bool(attr & 0x02)
                underline = bool(attr & 0x04)
            if len(pl) >= 56:
                cref = struct.unpack_from("<I", pl, 52)[0]
                r = cref & 0xFF
                g = (cref >> 8) & 0xFF
                b = (cref >> 16) & 0xFF
                # 검정·흰색 제외 (밝기 10~200 범위만 적용, 흰 배경에서 안 보이는 색 방지)
                lum = 0.299 * r + 0.587 * g + 0.114 * b
                if 10 < lum < 200:
                    color = f"#{r:02x}{g:02x}{b:02x}"
            charshape_map[idx] = {"bold": bold, "underline": underline, "color": color}
            idx += 1
    return charshape_map


def _parse_docinfo_parashapes(ole, compressed):
    """OLE2 DocInfo 스트림에서 ParaShape 목록 파싱.
    반환: {shape_idx(int): alignment_int} dict
    ParaShape 구조 (HWP5):
      Payload[0-3]: UINT32 속성
        bit0-1: alignment  0=양쪽, 1=왼쪽, 2=오른쪽, 3=가운데
    """
    docinfo_entry = ole.find_entry("docinfo")
    if not docinfo_entry:
        return {}
    raw = ole.read_stream(docinfo_entry)
    if not raw:
        return {}
    if compressed:
        for wbits in (-15, 15, 15 + 32):
            try:
                raw = zlib.decompress(raw, wbits)
                break
            except zlib.error:
                continue
    records = _parse_records(raw)
    parashape_map = {}
    idx = 0
    for rec in records:
        if rec["tag"] == HWPTAG_DOCINFO_PARA_SHAPE:
            pl = rec["payload"]
            alignment = 0
            if len(pl) >= 4:
                attr = struct.unpack_from("<I", pl, 0)[0]
                alignment = attr & 0x03  # bit0-1
            parashape_map[idx] = alignment
            idx += 1
    return parashape_map


def _extract_formatted_runs(text_payload, char_shapes, charshape_map, bold_texts, underline_texts):
    """단락 텍스트 + char_shapes + charshape_map → bold/underline 텍스트 추출."""
    if not char_shapes or not charshape_map:
        return
    # 텍스트를 (char_pos, char) 리스트로 추출
    chars_with_pos = []
    pos = 0
    char_pos = 0
    while pos + 2 <= len(text_payload):
        ch = struct.unpack_from("<H", text_payload, pos)[0]
        if ch >= 32:
            chars_with_pos.append((char_pos, chr(ch)))
            pos += 2
            char_pos += 1
        elif ch in _CTRL_16BYTE:
            pos += 16
            char_pos += 1
        else:
            pos += 2
            char_pos += 1
    if not chars_with_pos:
        return
    shapes_sorted = sorted(char_shapes, key=lambda x: x[0])

    def get_shape(cpos):
        sid = shapes_sorted[0][1] if shapes_sorted else 0
        for sp, s in shapes_sorted:
            if cpos >= sp:
                sid = s
            else:
                break
        return charshape_map.get(sid, {"bold": False, "underline": False})

    current_run = []
    current_bold = False
    current_underline = False

    def flush():
        t = "".join(current_run).strip()
        if t and 1 < len(t) <= 80:
            if current_bold:
                bold_texts.append(t)
            if current_underline:
                underline_texts.append(t)

    for cpos, ch in chars_with_pos:
        shape = get_shape(cpos)
        b, u = shape["bold"], shape["underline"]
        if b != current_bold or u != current_underline:
            flush()
            current_run = [ch]
            current_bold = b
            current_underline = u
        else:
            current_run.append(ch)
    flush()


def _collect_section_formatting(records, charshape_map, bold_texts, underline_texts):
    """섹션 레코드에서 PARA_TEXT + PARA_CHAR_SHAPE 쌍을 찾아 bold/underline 추출."""
    if not charshape_map:
        return
    current_text_payload = None
    current_char_shapes = []
    for rec in records:
        if rec["tag"] == HWPTAG_PARA_HEADER:
            if current_text_payload is not None:
                _extract_formatted_runs(
                    current_text_payload, current_char_shapes, charshape_map,
                    bold_texts, underline_texts
                )
            current_text_payload = None
            current_char_shapes = []
        elif rec["tag"] == HWPTAG_PARA_TEXT:
            current_text_payload = rec["payload"]
        elif rec["tag"] == HWPTAG_PARA_CHAR_SHAPE:
            shapes = []
            pl = rec["payload"]
            for i in range(0, len(pl) - 7, 8):
                sp = struct.unpack_from("<I", pl, i)[0]
                sid = struct.unpack_from("<I", pl, i + 4)[0]
                shapes.append((sp, sid))
            current_char_shapes = shapes
    # 마지막 단락 처리
    if current_text_payload is not None:
        _extract_formatted_runs(
            current_text_payload, current_char_shapes, charshape_map,
            bold_texts, underline_texts
        )


def parse_hwp5_native(data):
    """HWP 5.0 파일을 Python 표준 라이브러리만으로 파싱하여 HTML + text 반환.
    Returns: {'html': str, 'text': str, 'metadata': dict} or None
    """
    if data[:8] != OLE2_SIGNATURE:
        return None

    try:
        ole = OLE2Reader(data)
    except Exception as e:
        print(f"[HWP5] OLE2 파싱 실패: {e}")
        return None

    # FileHeader 확인
    fh_entry = ole.find_entry("fileheader")
    if not fh_entry:
        return None

    fh_data = ole.read_stream(fh_entry)
    if len(fh_data) < 40 or fh_data[:17] != b"HWP Document File":
        return None

    rev, build, minor, major = struct.unpack_from("BBBB", fh_data, 32)
    version = f"{major}.{minor}.{build}.{rev}"
    props = struct.unpack_from("<I", fh_data, 36)[0]
    compressed = bool(props & 0x01)
    encrypted = bool(props & 0x02)

    if encrypted:
        return {
            "html": "<p>암호화된 HWP 파일은 지원하지 않습니다.</p>",
            "text": "암호화된 HWP 파일은 지원하지 않습니다.",
            "metadata": {
                "warnings": ["HWP 파일이 암호화되어 있습니다."],
                "hwp_version": version,
            },
        }

    # BodyText/Section 스트림 수집
    section_entries = []
    for path, entry in ole.list_entries():
        if "bodytext" in path.lower() and "section" in entry["name"].lower():
            section_entries.append((path, entry))
    section_entries.sort(key=lambda x: x[0])

    if not section_entries:
        return None

    # 각 Section 파싱
    all_paragraphs = (
        []
    )  # list of {'type':'text','text':str} or {'type':'table','cells':[...]}
    warnings = []
    bold_texts = []
    underline_texts = []

    # DocInfo에서 CharShape (bold/underline) 맵 파싱
    charshape_map = _parse_docinfo_charshapes(ole, compressed)
    if charshape_map:
        pass  # charshape_map 파싱 완료

    # DocInfo에서 ParaShape (text-align) 맵 파싱
    parashape_map = _parse_docinfo_parashapes(ole, compressed)

    for sec_path, sec_entry in section_entries:
        raw = ole.read_stream(sec_entry)
        if not raw:
            continue

        # 압축 해제
        if compressed:
            try:
                raw = zlib.decompress(raw, -15)
            except zlib.error:
                try:
                    raw = zlib.decompress(raw)
                except zlib.error:
                    try:
                        raw = zlib.decompress(raw, 15 + 32)
                    except zlib.error as e:
                        warnings.append(f"{sec_path} 압축 해제 실패: {e}")
                        continue

        # 레코드 파싱
        records = _parse_records(raw)

        # 텍스트 + 테이블 구조 추출 (charshape_map, parashape_map 전달 → 서식 HTML 생성)
        _extract_content(records, all_paragraphs, charshape_map=charshape_map, parashape_map=parashape_map)

        # bold/underline 텍스트 추출
        _collect_section_formatting(records, charshape_map, bold_texts, underline_texts)

    if not all_paragraphs:
        return None

    # HTML + text 생성
    html_parts = []
    text_parts = []

    for para in all_paragraphs:
        if para["type"] == "text":
            txt = para["text"].strip()
            if txt:
                text_parts.append(txt)
                para_html = para.get("html", html_escape(txt))
                align = para.get("align", "")
                if align:
                    html_parts.append(f'<p style="text-align:{align}">{para_html}</p>')
                else:
                    html_parts.append(f"<p>{para_html}</p>")

        elif para["type"] == "table":
            cells = para["cells"]
            rows = para["rows"]
            cols = para["cols"]

            # ki-it.com 논문 방법론: lettered/unlettered 분류
            # 1행 테이블 또는 셀 수 ≤1 → unlettered (레이아웃 표 → <p> 변환)
            total_cells = len(cells)
            non_empty_cells = sum(
                1 for c in cells if any(t.strip() for t in c.get("texts", []))
            )
            fill_ratio = non_empty_cells / total_cells if total_cells > 0 else 0
            is_unlettered = (total_cells <= 1) or (rows == 1) or (fill_ratio < 0.15)

            if is_unlettered:
                # 레이아웃 표: 셀 텍스트를 <p>로 변환
                sorted_cells_u = sorted(cells, key=lambda c: (c["row"], c["col"]))
                for cell in sorted_cells_u:
                    cell_text = "\n".join(cell["texts"]) if cell["texts"] else ""
                    if cell_text.strip():
                        text_parts.append(cell_text)
                        cell_htmls = cell.get("htmls", [])
                        if cell_htmls:
                            html_parts.append(f"<p>{'<br>'.join(cell_htmls)}</p>")
                        else:
                            html_parts.append(f"<p>{html_escape(cell_text)}</p>")
                continue

            # occupied 셀 추적 (rowspan/colspan 처리)
            occupied = set()
            # 셀을 row, col 순으로 정렬
            sorted_cells = sorted(cells, key=lambda c: (c["row"], c["col"]))

            html_parts.append("<table>")
            for r in range(rows):
                html_parts.append("  <tr>")
                for cell in sorted_cells:
                    if cell["row"] != r:
                        continue
                    c = cell["col"]
                    if (r, c) in occupied:
                        continue

                    cell_text = "\n".join(cell["texts"]) if cell["texts"] else ""
                    text_parts.append(cell_text)
                    cell_tag = "th" if r == 0 else "td"
                    attrs = ""
                    rs = cell.get("rowspan", 1)
                    cs = cell.get("colspan", 1)
                    if rs > 1:
                        attrs += f' rowspan="{rs}"'
                    if cs > 1:
                        attrs += f' colspan="{cs}"'
                    cell_align = cell.get("align", "")
                    if cell_align:
                        attrs += f' style="text-align:{cell_align}"'
                    # 셀 영역 occupied 등록
                    for dr in range(rs):
                        for dc in range(cs):
                            occupied.add((r + dr, c + dc))

                    cell_htmls = cell.get("htmls", [])
                    if cell_htmls:
                        escaped = "<br>".join(cell_htmls)
                    else:
                        escaped = html_escape(cell_text).replace("\n", "<br>")
                    # 중첩 테이블 렌더링
                    nested_html = ""
                    for nested in cell.get("nested_tables", []):
                        if nested.get("type") == "table":
                            nc = nested.get("cells", [])
                            nr = nested.get("rows", 0)
                            sorted_nc = sorted(nc, key=lambda x: (x["row"], x["col"]))
                            nested_html += "<table>"
                            for nr_idx in range(nr):
                                nested_html += "<tr>"
                                for ncell in sorted_nc:
                                    if ncell["row"] != nr_idx:
                                        continue
                                    ntag = "th" if nr_idx == 0 else "td"
                                    ntxt = "\n".join(ncell.get("texts", []))
                                    nattrs = ""
                                    nrs = ncell.get("rowspan", 1)
                                    ncs = ncell.get("colspan", 1)
                                    if nrs > 1:
                                        nattrs += f' rowspan="{nrs}"'
                                    if ncs > 1:
                                        nattrs += f' colspan="{ncs}"'
                                    nalign = ncell.get("align", "")
                                    if nalign:
                                        nattrs += f' style="text-align:{nalign}"'
                                    nhtmls = ncell.get("htmls", [])
                                    if nhtmls:
                                        nesc = "<br>".join(nhtmls)
                                    else:
                                        nesc = html_escape(ntxt).replace("\n", "<br>")
                                    nested_html += f"<{ntag}{nattrs}>{nesc}</{ntag}>"
                                nested_html += "</tr>"
                            nested_html += "</table>"
                    html_parts.append(
                        f"    <{cell_tag}{attrs}>{escaped}{nested_html}</{cell_tag}>"
                    )
                html_parts.append("  </tr>")
            html_parts.append("</table>")

    html = "\n".join(html_parts)
    text = "\n".join(t for t in text_parts if t.strip())

    return {
        "html": html,
        "text": text,
        "metadata": {
            "hwp_version": version,
            "compressed": compressed,
            "sections": len(section_entries),
            "paragraphs": len(all_paragraphs),
            "warnings": warnings,
            "conversion_path": "hwp5 → native ole2 parser → html [b79]",
            "bold_texts": list(dict.fromkeys(bold_texts)),
            "underline_texts": list(dict.fromkeys(underline_texts)),
            "bullet_items": [],
        },
    }


def _parse_records(stream_data):
    """HWP 바이너리 레코드 스트림을 파싱."""
    records = []
    pos = 0
    while pos + 4 <= len(stream_data):
        header = struct.unpack_from("<I", stream_data, pos)[0]
        tag_id = header & 0x3FF
        level = (header >> 10) & 0x3FF
        size = (header >> 20) & 0xFFF
        pos += 4

        if size == 0xFFF and pos + 4 <= len(stream_data):
            size = struct.unpack_from("<I", stream_data, pos)[0]
            pos += 4

        payload = stream_data[pos : pos + size]
        pos += size

        records.append(
            {
                "tag": tag_id,
                "level": level,
                "payload": payload,
            }
        )
    return records


def _extract_text_from_para(payload):
    """HWPTAG_PARA_TEXT 페이로드에서 텍스트 추출 (UTF-16LE + 제어문자 처리)."""
    parts = []
    pos = 0
    while pos + 2 <= len(payload):
        ch = struct.unpack_from("<H", payload, pos)[0]
        if ch >= 32:
            parts.append(chr(ch))
            pos += 2
        elif ch == 13 or ch == 10:
            parts.append("\n")
            pos += 2
        elif ch == 30:
            parts.append("-")
            pos += 2
        elif ch in _CTRL_16BYTE:
            pos += 16
            if pos > len(payload):
                break
        else:
            pos += 2
    return "".join(parts)


def _build_formatted_html(text_payload, char_shapes, charshape_map):
    """PARA_TEXT payload + PARA_CHAR_SHAPE + charshape_map → 서식 적용 HTML.
    <strong>(볼드), <u>(밑줄) 태그 인라인 적용. 텍스트는 html_escape 처리.
    char_shapes가 없으면 전체 텍스트를 html_escape만 적용하여 반환.
    """
    # 텍스트를 (char_pos, char) 리스트로 추출 (UTF-16LE + 제어문자 처리)
    chars_with_pos = []
    pos = 0
    char_pos = 0
    while pos + 2 <= len(text_payload):
        ch = struct.unpack_from("<H", text_payload, pos)[0]
        if ch >= 32:
            chars_with_pos.append((char_pos, chr(ch)))
            pos += 2
            char_pos += 1
        elif ch == 13 or ch == 10:
            chars_with_pos.append((char_pos, "\n"))
            pos += 2
            char_pos += 1
        elif ch == 30:
            chars_with_pos.append((char_pos, "-"))
            pos += 2
            char_pos += 1
        elif ch in _CTRL_16BYTE:
            pos += 16
            char_pos += 1
            if pos > len(text_payload):
                break
        else:
            pos += 2
            char_pos += 1

    if not chars_with_pos:
        return ""

    # char_shapes가 없거나 charshape_map이 비어있으면 plain html_escape 반환
    plain_text = "".join(c for _, c in chars_with_pos).rstrip("\n")
    if not char_shapes or not charshape_map:
        return html_escape(plain_text).replace("\n", "<br>")

    shapes_sorted = sorted(char_shapes, key=lambda x: x[0])

    def get_shape(cpos):
        sid = shapes_sorted[0][1] if shapes_sorted else 0
        for sp, s in shapes_sorted:
            if cpos >= sp:
                sid = s
            else:
                break
        return charshape_map.get(sid, {"bold": False, "underline": False, "color": None})

    # 같은 서식의 연속 문자를 run으로 묶기
    runs = []  # list of (bold, underline, color, text_str)
    current_run_chars = []
    current_bold = None
    current_underline = None
    current_color = None

    for cpos, ch in chars_with_pos:
        shape = get_shape(cpos)
        b, u, c = shape["bold"], shape["underline"], shape.get("color")
        if b != current_bold or u != current_underline or c != current_color:
            if current_run_chars:
                runs.append((current_bold, current_underline, current_color, "".join(current_run_chars)))
            current_run_chars = [ch]
            current_bold = b
            current_underline = u
            current_color = c
        else:
            current_run_chars.append(ch)
    if current_run_chars:
        runs.append((current_bold, current_underline, current_color, "".join(current_run_chars)))

    # 후행 개행 제거
    if runs:
        last_b, last_u, last_c, last_text = runs[-1]
        runs[-1] = (last_b, last_u, last_c, last_text.rstrip("\n"))

    # 각 run을 HTML로 변환
    html_out = []
    for bold, underline, color, text in runs:
        if not text:
            continue
        escaped = html_escape(text).replace("\n", "<br>")
        if bold:
            escaped = f"<strong>{escaped}</strong>"
        if underline:
            escaped = f"<u>{escaped}</u>"
        if color:
            escaped = f'<span style="color:{color}">{escaped}</span>'
        html_out.append(escaped)

    return "".join(html_out)


def _extract_content(records, all_paragraphs, charshape_map=None, parashape_map=None):
    """레코드 리스트에서 텍스트와 테이블 구조를 추출 (중첩 테이블 스택 지원)."""
    table_stack = []  # 중첩 테이블 컨텍스트 스택
    current_table = None
    current_cell = None
    table_cells = []
    table_start_level = -1
    awaiting_para_text = False
    # 서식 적용: PARA_TEXT → PARA_CHAR_SHAPE 순서이므로 지연 적용
    current_text_payload = None  # PARA_TEXT raw payload (서식 생성용)
    last_fmt_ref = None  # ('cell', cell_dict) or ('para', para_dict)
    current_para_align = ""  # 현재 단락 정렬 (PARA_HEADER에서 읽어 PARA_TEXT에 적용)

    def _build_table_para():
        if not table_cells:
            return None
        rows = current_table.get("rows", 0) if current_table else 0
        cols = current_table.get("cols", 0) if current_table else 0
        if table_cells:
            max_row = max(c["row"] + c["rowspan"] for c in table_cells)
            max_col = max(c["col"] + c["colspan"] for c in table_cells)
            rows = max(rows, max_row)
            cols = max(cols, max_col)
        return {"type": "table", "rows": rows, "cols": cols, "cells": table_cells}

    def _pop_nested():
        nonlocal current_table, table_cells, current_cell, table_start_level
        nested_para = _build_table_para()
        parent = table_stack.pop()
        current_table = parent["table"]
        table_cells = parent["cells"]
        current_cell = parent["cell"]
        table_start_level = parent["start_level"]
        # 중첩 테이블을 부모 셀에 추가
        if parent["cell"] is not None and nested_para:
            parent["cell"].setdefault("nested_tables", []).append(nested_para)
        elif nested_para:
            all_paragraphs.append(nested_para)

    for rec in records:
        # 레벨 기반 자동 pop: 현재 셀이 없고 레코드 레벨이 테이블 시작 레벨 이하이면 pop
        while (
            table_stack
            and current_cell is None
            and current_table
            and table_cells
            and rec["level"] <= table_start_level
        ):
            _pop_nested()
            awaiting_para_text = False

        # 빈 문단 감지: PARA_HEADER로 셀 내 문단 수 정확히 추적
        if rec["tag"] == HWPTAG_PARA_HEADER:
            current_text_payload = None
            last_fmt_ref = None
            # ParaShape에서 정렬 정보 읽기 (payload offset 12~13: para_shape_id WORD)
            current_para_align = ""
            if parashape_map:
                pl = rec["payload"]
                if len(pl) >= 14:
                    shape_id = struct.unpack_from("<H", pl, 12)[0]
                    alignment = parashape_map.get(shape_id, 0)
                    if alignment == 2:
                        current_para_align = "right"
                    elif alignment == 3:
                        current_para_align = "center"
            if current_cell is not None and awaiting_para_text:
                current_cell["para_count"] += 1
                if current_cell["para_count"] >= current_cell["para_target"]:
                    current_cell = None
            awaiting_para_text = current_cell is not None

        if rec["tag"] == HWPTAG_CTRL_HEADER and len(rec["payload"]) >= 4:
            ctrl_id = rec["payload"][:4]
            if ctrl_id in (b"tbl ", b" lbt"):
                if current_table:
                    if current_cell is not None and awaiting_para_text:
                        current_cell["para_count"] += 1
                    # 셀이 모두 완료된 테이블 → flush 후 새 테이블 시작 (순서 보장)
                    if current_cell is None and table_cells:
                        tbl_para = _build_table_para()
                        if tbl_para:
                            all_paragraphs.append(tbl_para)
                        table_cells = []
                    else:
                        # 진행 중인 셀 있음 → 진짜 중첩: 스택에 push
                        table_stack.append(
                            {
                                "table": current_table,
                                "cells": table_cells,
                                "cell": current_cell,
                                "start_level": table_start_level,
                            }
                        )
                        table_cells = []
                current_table = {"rows": 0, "cols": 0}
                table_start_level = rec["level"]
                current_cell = None
                awaiting_para_text = False

        elif rec["tag"] == HWPTAG_TABLE and current_table and len(rec["payload"]) >= 8:
            rows = struct.unpack_from("<H", rec["payload"], 4)[0]
            cols = struct.unpack_from("<H", rec["payload"], 6)[0]
            current_table["rows"] = rows
            current_table["cols"] = cols

            # 루트 레벨 1x1 레이아웃 래퍼 테이블 해소
            # 문서 전체가 1x1 셀 안에 감싸진 HWP 구조 → 래퍼를 제거하여
            # 내부 콘텐츠 테이블이 all_paragraphs에 직접 추가되도록 함
            if rows == 1 and cols == 1 and not table_stack:
                current_table = None
                table_cells = []
                current_cell = None
                table_start_level = -1
                awaiting_para_text = False

        elif rec["tag"] == HWPTAG_LIST_HEADER and current_table:
            if current_cell is not None and awaiting_para_text:
                current_cell["para_count"] += 1
            current_cell = None
            awaiting_para_text = False
            payload = rec["payload"]
            if len(payload) >= 16:
                num_para = struct.unpack_from("<H", payload, 0)[0]
                col_addr = struct.unpack_from("<H", payload, 8)[0]
                row_addr = struct.unpack_from("<H", payload, 10)[0]
                col_span = struct.unpack_from("<H", payload, 12)[0]
                row_span = struct.unpack_from("<H", payload, 14)[0]
                current_cell = {
                    "row": row_addr,
                    "col": col_addr,
                    "rowspan": max(row_span, 1),
                    "colspan": max(col_span, 1),
                    "texts": [],
                    "nested_tables": [],
                    "para_target": num_para,
                    "para_count": 0,
                }
                table_cells.append(current_cell)

        elif rec["tag"] == HWPTAG_PARA_TEXT:
            awaiting_para_text = False
            text = _extract_text_from_para(rec["payload"]).rstrip("\n")
            current_text_payload = rec["payload"]  # PARA_CHAR_SHAPE에서 서식 생성용
            last_fmt_ref = None

            if current_cell is not None:
                if text.strip():
                    current_cell["texts"].append(text)
                    last_fmt_ref = ("cell", current_cell)
                    # 셀의 첫 번째 단락 정렬을 셀 정렬로 저장
                    if "align" not in current_cell and current_para_align:
                        current_cell["align"] = current_para_align
                current_cell["para_count"] += 1
                if current_cell["para_count"] >= current_cell["para_target"]:
                    current_cell = None
            elif current_table and table_cells and table_stack:
                _pop_nested()
                if current_cell is not None:
                    if text.strip():
                        current_cell["texts"].append(text)
                        last_fmt_ref = ("cell", current_cell)
                        if "align" not in current_cell and current_para_align:
                            current_cell["align"] = current_para_align
                    current_cell["para_count"] += 1
                    if current_cell["para_count"] >= current_cell["para_target"]:
                        current_cell = None
                elif text.strip():
                    para_dict = {"type": "text", "text": text, "align": current_para_align}
                    all_paragraphs.append(para_dict)
                    last_fmt_ref = ("para", para_dict)
            else:
                if current_table and table_cells:
                    tbl_para = _build_table_para()
                    if tbl_para:
                        all_paragraphs.append(tbl_para)
                    table_cells = []
                    current_table = None
                    current_cell = None
                if text.strip():
                    para_dict = {"type": "text", "text": text, "align": current_para_align}
                    all_paragraphs.append(para_dict)
                    last_fmt_ref = ("para", para_dict)

        elif rec["tag"] == HWPTAG_PARA_CHAR_SHAPE:
            # PARA_TEXT 뒤에 오는 서식 정보 → 지연 적용
            shapes = []
            pl = rec["payload"]
            for i in range(0, len(pl) - 7, 8):
                sp = struct.unpack_from("<I", pl, i)[0]
                sid = struct.unpack_from("<I", pl, i + 4)[0]
                shapes.append((sp, sid))
            if current_text_payload is not None and shapes and charshape_map:
                fmt_html = _build_formatted_html(current_text_payload, shapes, charshape_map)
                if last_fmt_ref:
                    kind, target = last_fmt_ref
                    if kind == "cell":
                        target.setdefault("htmls", []).append(fmt_html)
                    elif kind == "para":
                        target["html"] = fmt_html
                last_fmt_ref = None
                current_text_payload = None

    # 마지막 빈 문단 처리
    if current_cell is not None and awaiting_para_text:
        current_cell["para_count"] += 1
        if current_cell["para_count"] >= current_cell["para_target"]:
            current_cell = None

    # 잔여 스택 flush
    while table_stack:
        _pop_nested()
    if current_table and table_cells:
        tbl_para = _build_table_para()
        if tbl_para:
            all_paragraphs.append(tbl_para)


# ============================================
# Fallback: stdlib-only HTTP Server
# ============================================


def parse_multipart(fp, content_type):
    """cgi.FieldStorage 대체: multipart/form-data에서 파일과 텍스트 필드 추출"""
    boundary = None
    for part in content_type.split(";"):
        part = part.strip()
        if part.startswith("boundary="):
            boundary = part[len("boundary=") :].strip('"')
            break
    if not boundary:
        return None, None, ""

    raw = fp.read()
    boundary_bytes = ("--" + boundary).encode()
    parts = raw.split(boundary_bytes)

    filename = None
    file_data = None
    api_key = ""

    for part in parts:
        if b"Content-Disposition" not in part:
            continue
        # 헤더와 바디 분리
        header_end = part.find(b"\r\n\r\n")
        if header_end < 0:
            continue
        headers_raw = part[:header_end].decode("utf-8", errors="replace")
        body = part[header_end + 4 :]
        # 후행 \r\n 제거
        if body.endswith(b"\r\n"):
            body = body[:-2]
        if body.endswith(b"--\r\n"):
            body = body[:-4]
        if body.endswith(b"--"):
            body = body[:-2]

        # anthropic_api_key 텍스트 필드
        if 'name="anthropic_api_key"' in headers_raw:
            api_key = body.decode("utf-8", errors="replace").strip()
            continue

        # name="file" 확인
        if 'name="file"' not in headers_raw:
            continue

        # filename 추출
        fn_match = re.search(r'filename="([^"]+)"', headers_raw)
        filename = fn_match.group(1) if fn_match else "unknown"
        file_data = body

    if filename and file_data is not None:
        return filename, file_data, api_key
    return None, None, ""


def run_stdlib_server(port):
    """Flask가 없을 때 표준 라이브러리로 서버 실행"""
    from http.server import HTTPServer, BaseHTTPRequestHandler

    class ConvertHandler(BaseHTTPRequestHandler):
        def _cors(self):
            self.send_header("Access-Control-Allow-Origin", "*")
            self.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
            self.send_header("Access-Control-Allow-Headers", "*")

        def do_OPTIONS(self):
            self.send_response(204)
            self._cors()
            self.end_headers()

        def do_GET(self):
            if self.path == "/health":
                self.send_response(200)
                self._cors()
                self.send_header("Content-Type", "application/json")
                self.end_headers()
                resp = json.dumps(
                    {
                        "status": "ok",
                        "version": "build75",
                        "formats": get_supported_formats(),
                        "libreoffice": SOFFICE_PATH or False,
                        "hwp_table_support": bool(SOFFICE_PATH),
                        "playwright": check_playwright(),
                    }
                )
                self.wfile.write(resp.encode())
            else:
                self.send_response(404)
                self.end_headers()

        def do_POST(self):
            allowed = ("/api/convert", "/api/hwp-to-pdf", "/api/hwp-to-rawhtml", "/api/render-kv", "/api/render-full")
            if not any(self.path.startswith(p) for p in allowed):
                self.send_response(404)
                self.end_headers()
                return

            # /api/render-kv, /api/render-full는 JSON body
            if self.path.startswith("/api/render-kv") or self.path.startswith("/api/render-full"):
                content_length = int(self.headers.get("Content-Length", 0))
                body = self.rfile.read(content_length)
                try:
                    payload = json.loads(body)
                    if self.path.startswith("/api/render-full"):
                        img_bytes, mime, err = _render_full_playwright(payload)
                    else:
                        img_bytes, mime, err = _render_kv_playwright(payload)
                    if err:
                        self.send_response(500)
                        self._cors()
                        self.send_header("Content-Type", "application/json")
                        self.end_headers()
                        self.wfile.write(json.dumps({"error": err}).encode())
                    else:
                        self.send_response(200)
                        self._cors()
                        self.send_header("Content-Type", mime)
                        self.send_header("Content-Length", str(len(img_bytes)))
                        self.end_headers()
                        self.wfile.write(img_bytes)
                except Exception as e:
                    self.send_response(500)
                    self._cors()
                    self.send_header("Content-Type", "application/json")
                    self.end_headers()
                    self.wfile.write(json.dumps({"error": str(e)}).encode())
                return

            content_type = self.headers.get("Content-Type", "")
            if "multipart/form-data" not in content_type:
                self.send_response(400)
                self._cors()
                self.send_header("Content-Type", "application/json")
                self.end_headers()
                self.wfile.write(
                    json.dumps({"error": "multipart/form-data required"}).encode()
                )
                return

            content_length = int(self.headers.get("Content-Length", 0))
            body_stream = BytesIO(self.rfile.read(content_length))

            filename, data, form_api_key = parse_multipart(body_stream, content_type)
            if not filename or not data:
                self.send_response(400)
                self._cors()
                self.end_headers()
                self.wfile.write(json.dumps({"error": "No file"}).encode())
                return

            ext = filename.rsplit(".", 1)[-1].lower()

            try:
                if self.path.startswith("/api/hwp-to-rawhtml"):
                    html_out, warns = convert_to_raw_html(data, ext)
                    if html_out:
                        result = {"html": html_out, "warnings": warns}
                        self.send_response(200)
                    else:
                        result = {"error": (warns[0] if warns else "변환 실패")}
                        self.send_response(500)
                elif self.path.startswith("/api/hwp-to-pdf"):
                    import base64

                    pdf_data, err = convert_any_to_pdf(data, ext, filename)
                    if pdf_data:
                        result = {
                            "pdf": base64.b64encode(pdf_data).decode(),
                            "warnings": [],
                        }
                        self.send_response(200)
                    else:
                        result = {"error": err or "변환 실패"}
                        self.send_response(500)
                else:
                    from urllib.parse import urlparse, parse_qs
                    parsed = urlparse(self.path)
                    parser_hint = parse_qs(parsed.query).get("parser", [""])[0]
                    api_key = form_api_key or os.environ.get('ANTHROPIC_API_KEY', '')
                    result = convert_file(data, ext, filename, parser_hint=parser_hint, api_key=api_key)
                    self.send_response(200)
            except Exception as e:
                result = {
                    "error": str(e),
                    "html": "",
                    "text": "",
                    "metadata": {"warnings": [str(e)]},
                }
                self.send_response(500)

            self._cors()
            self.send_header("Content-Type", "application/json")
            self.end_headers()
            self.wfile.write(json.dumps(result, ensure_ascii=False).encode())

        def log_message(self, fmt, *args):
            print(f"[Convert] {args[0] if args else ''}")

    server = HTTPServer(("0.0.0.0", port), ConvertHandler)
    print(f"Document Convert Server (stdlib) on http://localhost:{port}")
    print(
        f"  POST /api/convert  |  POST /api/hwp-to-rawhtml  |  POST /api/hwp-to-pdf  |  GET /health"
    )
    print("Press Ctrl+C to stop")
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print("\nStopping...")
        server.server_close()


# ============================================
# Main
# ============================================


def main():
    port = int(sys.argv[1]) if len(sys.argv) > 1 else 8082

    if FLASK_AVAILABLE:
        print(f"Document Convert Server (Flask) on http://localhost:{port}")
        print(f"  POST /api/convert  |  GET /health")
        print(f"Supported formats: {get_supported_formats()}")
        app.run(host="0.0.0.0", port=port, debug=False)
    else:
        print("Flask not found — using stdlib HTTP server")
        print("For better performance: pip3 install flask")
        run_stdlib_server(port)


if __name__ == "__main__":
    main()
